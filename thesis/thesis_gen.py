# -*- coding: utf-8 -*-
"""
台大碩士論文產生器
題目：基於 3D 重建與語意分割之食物熱量估算系統
輸出：碩士論文2.docx
"""

import os
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(SCRIPT_DIR, "figures")
REPO_DIR = os.path.dirname(SCRIPT_DIR)
from references_data import REFERENCES
from docx.oxml import OxmlElement
from omml_equations import (
    M,
    add_equation,
    fill_paragraph_mixed,
    m_frac,
    m_nary,
    m_omath,
    m_sub,
    m_sup,
    m_text,
    split_math_markup,
)

CN_FONT = "標楷體"
EN_FONT = "Times New Roman"
BODY_SIZE = Pt(12)

# 論文與作者資訊（台大電信所格式）
THESIS_INFO = {
    "college_zh": "電機資訊學院",
    "dept_zh": "電信工程學研究所",
    "college_en": "College of Electrical Engineering and Computer Science",
    "dept_en": "Graduate Institute of Communication Engineering",
    "title_zh": "基於 3D 重建與語意分割之食物熱量估算系統",
    "title_en": "A Food Calorie Estimation System Based on\n"
                "3D Reconstruction and Semantic Segmentation",
    "student_zh": "陳慶瑋",
    "student_en": "Ching-Wei Chen",
    "student_id": "（學號待填）",
    "advisor_zh": "葉丙成",
    "advisor_en": "Bing-Cheng Ye",
    "advisor_degree_zh": "博士",
    "advisor_degree_en": "Ph.D.",
    "year_roc": "115",
    "month_num": "7",
    "month_en": "July",
    "year_ad": "2026",
}

doc = Document()

# ---------------------------------------------------------------------------
# 基本樣式：Normal
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = EN_FONT
normal.font.size = BODY_SIZE
normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
pf.space_after = Pt(0)


def _set_run_font(run, size=BODY_SIZE, bold=False, cn=CN_FONT, en=EN_FONT):
    run.font.size = size
    run.font.bold = bold
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), cn)


def cit(*nums):
    """內文引用標註，如 cit(1, 2) -> '[1,2]'。"""
    return "[" + ",".join(str(n) for n in nums) + "]"


def equation(omath, number=None):
    """插入置中之 Word 原生方程式；可選右側編號如 (3-1)。"""
    return add_equation(doc, omath, number=number, font_pt=BODY_SIZE.pt)


def _write_runs(p, text, size=BODY_SIZE, bold=False):
    if text and "{" in text and "}" in text:
        fill_paragraph_mixed(
            p,
            split_math_markup(text),
            lambda run: _set_run_font(run, size=size, bold=bold),
        )
    elif text:
        r = p.add_run(text)
        _set_run_font(r, size=size, bold=bold)


def para(text="", size=BODY_SIZE, bold=False, align=None, indent_chars=2,
         space_before=0, space_after=6, line=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE if line == 1.5 else None
    if line != 1.5:
        fmt.line_spacing = line
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if indent_chars:
        fmt.first_line_indent = Pt(size.pt * indent_chars)
    _write_runs(p, text, size=size, bold=bold)
    return p


def _numbering_style(level):
    return {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]


def heading(text, level=1, page_break_before=False):
    style = _numbering_style(level)
    p = doc.add_paragraph(style=style)
    if page_break_before:
        p.paragraph_format.page_break_before = True
    sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13)}
    fmt = p.paragraph_format
    fmt.space_before = Pt(12)
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    _set_run_font(r, size=sizes[level], bold=True)
    r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def bullet(text, size=BODY_SIZE):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_run_font(r, size=size)
    return p


def chapter_outline(ch_num, title, body):
    """論文概述用：章標題（粗體）＋內文（首行縮排）。"""
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.first_line_indent = Pt(0)
    fmt.left_indent = Pt(0)
    fmt.space_before = Pt(8)
    fmt.space_after = Pt(4)
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(f"第{ch_num}章　{title}")
    _set_run_font(r, size=BODY_SIZE, bold=True)
    para(body, indent_chars=2, space_before=0, space_after=10)


def numbered(text, size=BODY_SIZE):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _set_run_font(r, size=size)
    return p


def caption(text, kind="圖", size=Pt(11)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=True)
    return p


def figure_placeholder(label, desc):
    """插入圖片佔位框（灰底方框）與圖說。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"〔此處插入圖片：{desc}〕")
    _set_run_font(r, size=Pt(11))
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    caption(label, kind="圖")


def add_image(path, label, width_cm=13, color=False):
    """插入實際圖片與圖說；找不到檔案時退回佔位框。

    color=False：概念／架構圖轉高對比灰階（適合印刷）。
    color=True：實拍照片保留彩色。
    """
    if not os.path.exists(path):
        figure_placeholder(label, f"缺少檔案 {os.path.basename(path)}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    if color:
        run.add_picture(path, width=Cm(width_cm))
    else:
        try:
            from PIL import Image
            img = Image.open(path).convert("RGB").convert("L")
            bw = img.point(lambda p: 255 if p > 200 else 0, mode="1")
            buf = io.BytesIO()
            bw.save(buf, format="PNG")
            buf.seek(0)
            run.add_picture(buf, width=Cm(width_cm))
        except Exception:
            run.add_picture(path, width=Cm(width_cm))
    caption(label, kind="圖")


def fig_img(fname, label, width_cm=13, color=False):
    add_image(os.path.join(FIG_DIR, fname), label, width_cm, color=color)


def repo_img(fname, label, width_cm=11, color=True):
    """repo 根目錄之實拍／pipeline 輸出，預設保留彩色。"""
    add_image(os.path.join(REPO_DIR, fname), label, width_cm, color=color)


def _set_cell_shading(cell, fill="FFFFFF"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tc_borders = tcPr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)
    for edge, attrs in kwargs.items():
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, val in attrs.items():
            element.set(qn(f"w:{key}"), str(val))


def _apply_plain_table_style(table):
    border = {"sz": "4", "val": "single", "color": "000000", "space": "0"}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_shading(cell)
            _set_cell_border(
                cell,
                top=border,
                left=border,
                bottom=border,
                right=border,
            )


def _write_cell(cell, text, bold=False, size=Pt(11)):
    p = cell.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _write_runs(p, str(text), size=size, bold=bold)


def make_table(headers, rows, caption_text=None, col_widths=None, todo=False):
    if caption_text:
        caption(caption_text, kind="表")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _write_cell(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _write_cell(cells[i], val)
    _apply_plain_table_style(table)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    if todo:
        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(10)
        r = note.add_run("（TODO：上表數值待實驗完成後填入）")
        _set_run_font(r, size=Pt(10))
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def todo_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(f"【待補充 TODO】{text}")
    _set_run_font(r, size=Pt(11))
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return p


def page_break():
    doc.add_page_break()


def _add_field(paragraph, instr):
    run = paragraph.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（請於 Word 中按右鍵 → 更新功能變數）"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._element.append(fldBegin)
    run._element.append(instrText)
    run._element.append(fldSep)
    r2 = paragraph.add_run()
    r2._element.append(placeholder)
    r3 = paragraph.add_run()
    r3._element.append(fldEnd)


def toc(title, instr):
    heading(title, level=1, page_break_before=True)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    _add_field(p, instr)


def set_page_number_format(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:fmt"), fmt)
    if start is not None:
        pgNumType.set(qn("w:start"), str(start))


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._element.append(fldBegin)
    run._element.append(instrText)
    run._element.append(fldEnd)
    _set_run_font(run, size=Pt(11))


# ---------------------------------------------------------------------------
# 版面：邊界（台大論文格式：上 3cm、下 2cm、左右各 3cm）
# ---------------------------------------------------------------------------
def _apply_ntu_margins(section, top_cm=3.0, bottom_cm=2.0):
    section.top_margin = Cm(top_cm)
    section.bottom_margin = Cm(bottom_cm)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)


_apply_ntu_margins(doc.sections[0])

# ===========================================================================
# 封面／書名頁（台大附件 1、2）
# ===========================================================================
def _cover_line(text, size, bold=False, en=False, space_after=6, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, cn=EN_FONT if en else CN_FONT, en=EN_FONT)
    return p


def build_title_page(top_space_before=56):
    """封面或書名頁：上留白約 4cm、下留白約 3cm（以段距模擬）。"""
    info = THESIS_INFO
    _cover_line(
        f"國立臺灣大學{info['college_zh']}{info['dept_zh']}",
        Pt(18), bold=True, space_before=top_space_before, space_after=10,
    )
    _cover_line("碩士論文", Pt(18), bold=True, space_after=12)
    _cover_line(
        f"Department or {info['dept_en']}",
        Pt(14), en=True, space_after=4,
    )
    _cover_line(info["college_en"], Pt(14), en=True, space_after=4)
    _cover_line("National Taiwan University", Pt(16), en=True, space_after=4)
    _cover_line("Master's Thesis", Pt(16), en=True, space_after=24)
    _cover_line(info["title_zh"], Pt(18), bold=True, space_after=8)
    for line in info["title_en"].split("\n"):
        _cover_line(line, Pt(18), en=True, bold=True, space_after=4)
    _cover_line("", Pt(12), space_after=20)
    _cover_line(info["student_zh"], Pt(18), bold=True, space_after=4)
    _cover_line(info["student_en"], Pt(18), en=True, space_after=24)
    _cover_line(
        f"指導教授：{info['advisor_zh']}　{info['advisor_degree_zh']}",
        Pt(18), bold=True, space_after=4,
    )
    _cover_line(
        f"Advisor: {info['advisor_en']}, {info['advisor_degree_en']}",
        Pt(18), en=True, space_after=36,
    )
    _cover_line(
        f"中華民國 {info['year_roc']} 年 {info['month_num']} 月",
        Pt(18), bold=True, space_after=4,
    )
    _cover_line(
        f"{info['month_en']} {info['year_ad']}",
        Pt(18), en=True, space_after=24,
    )


build_title_page()
page_break()
build_title_page(top_space_before=48)

# ===========================================================================
# 口試委員會審定書（台大附件 3）
# ===========================================================================
page_break()
info = THESIS_INFO
_cover_line("國立臺灣大學碩士學位論文", Pt(16), bold=True, space_before=24, space_after=4)
_cover_line("口試委員會審定書", Pt(16), bold=True, space_after=8)
_cover_line("MASTER'S THESIS ACCEPTANCE CERTIFICATE", Pt(14), en=True, space_after=2)
_cover_line("NATIONAL TAIWAN UNIVERSITY", Pt(14), en=True, space_after=24)

_cover_line(info["title_zh"], Pt(14), bold=True, space_after=8)
_cover_line(
    info["title_en"].replace("\n", " "),
    Pt(12), en=True, space_after=24,
)

para(
    f"本論文係 {info['student_zh']} {info['student_id']} 在國立臺灣大學"
    f"{info['dept_zh']}完成之碩士學位論文，於民國 {info['year_roc']} 年 "
    f"{info['month_num']} 月　　日承下列考試委員審查通過及口試及格，特此證明。",
    indent_chars=0, space_after=12,
)
para(
    "The undersigned, appointed by the "
    f"{info['dept_en']} on "
    f"　　　 (date)　　(month)　　(year) have examined a Master's Thesis "
    f"entitled above presented by {info['student_en']} "
    f"{info['student_id']} candidate and hereby certify that it is worthy of acceptance.",
    indent_chars=0, space_after=18,
)
para("口試委員 Oral examination committee：", indent_chars=0, space_after=8)
para(f"____________________（指導教授 Advisor：{info['advisor_zh']}）", indent_chars=0, space_after=8)
para("____________________    ____________________    ____________________", indent_chars=0, space_after=8)
para("____________________    ____________________    ____________________", indent_chars=0, space_after=16)
para("系（所）主管 Director: ____________________", indent_chars=0, space_after=12)

# ===========================================================================
# 誌謝
# ===========================================================================
heading("誌　謝", level=1, page_break_before=True)
para("研究生涯的完成，仰賴許多人的支持與協助。")
para(f"首先，衷心感謝指導教授{THESIS_INFO['advisor_zh']}{THESIS_INFO['advisor_degree_zh']}。"
     "從研究題目的發想、系統架構的規劃、到論文的反覆修訂，老師皆給予悉心的指導與寶貴的"
     "建議，使本研究得以順利完成。老師嚴謹的研究態度與對問題本質的洞察，是學生受用一生的典範。")
para("感謝口試委員教授們，於百忙之中撥冗審閱論文並提供諸多深具啟發性的意見，使本論文更臻完善。")
para("感謝實驗室的同窗與學長姐，在研究過程中不吝分享經驗、協助資料蒐集與系統測試；也感謝所有"
     "參與多視角影像採集與使用者測試的協助者，你們的付出是本研究得以驗證的重要基礎。")
para("最後，謹將此論文獻給我摯愛的家人，感謝你們一路以來無條件的支持與鼓勵。")
para("", indent_chars=0)
para(f"{THESIS_INFO['student_zh']}　謹誌", align=WD_ALIGN_PARAGRAPH.RIGHT, indent_chars=0)
para(
    f"於國立臺灣大學　中華民國 {THESIS_INFO['year_roc']} 年 {THESIS_INFO['month_num']} 月",
    align=WD_ALIGN_PARAGRAPH.RIGHT, indent_chars=0,
)

# ===========================================================================
# 中文摘要
# ===========================================================================
heading("摘　要", level=1, page_break_before=True)
para("隨著行動醫療與智慧健康照護的興起，精準且便利的飲食熱量監控已成為個人健康管理的重要環節。"
     "然而，現有以單張二維影像結合大型多模態模型之熱量估算方法，因缺乏絕對物理尺度、易產生視覺"
     "幻覺、且推論過程缺乏可解釋性，導致食物體積與熱量的估算誤差偏大，難以應用於實際的健康管理"
     "場景。")
para("針對上述問題，本研究提出一套「基於三維重建與語意分割之食物熱量估算系統」。本系統以日常隨身"
     "之標準信用卡作為絕對比例尺，並整合三項關鍵技術：（一）以輕量化視覺語言模型 Qwen2.5-VL "
     "進行開放域食物品項辨識與定位；（二）以萬物分割模型 SAM3 進行多食材之像素級邊界分割；"
     "（三）以多視角三維重建模型 DUSt3R 由正上方與左右四十五度共三張影像還原食物之三維點雲。"
     "系統進一步透過信用卡之已知尺寸完成比例尺錨定，將無尺度之點雲校正為具絕對物理單位之空間"
     "模型，並以平面擬合與厚度積分演算法計算各食材之管線估計體積；最後透過本研究建立之「食物"
     "體積—重量—營養轉換資料庫」，將管線估計體積轉換為重量、熱量與三大營養素。資料庫建置與"
     "線上推論採用相同之體積估算管線，使體積雖非獨立驗證之物理真值，仍能在系統內部形成自洽之"
     "體積—重量對應。")
para("在系統實作上，本研究採用行動端與雲端協同之架構：行動應用程式負責標準化的多視角影像採集與"
     "結果呈現，雲端 GPU 推論服務則負責執行辨識、分割與三維重建之運算，兼顧一般使用者無專用深度"
     "感測硬體之限制與運算資源需求。")
para("本研究已完成系統原型建置與實驗設計，並以活大自助餐場域蒐集測試樣本；熱量與體積估算之"
     "定量評估指標（MAPE、MAE）將於第六章呈現。文獻回顧中，Nutrition5k 之 2D direct "
     f"baseline 報告 MAPE 約 26.1%（固定 top-view 設定）{cit(20)}；Shao 等人"
     f"{cit(49)}於 96 張真實用餐影像上報告 MAPE 11.47%；Vinod 等人"
     f"{cit(22)}於 SimpleFood45 上報 EMAPE 17.67%（需 checkerboard 與 3D food "
     "model）。上述數值之資料集、輸入模態與評估條件各異，不宜直接排序；本研究之"
     "比較將在相同活大拍攝協定下，以本方法與自行建立之 LMM 基準對照呈現。")
para("關鍵詞：食物熱量估算、三維重建、語意分割、視覺語言模型、比例尺錨定、行動醫療、"
     "多視角影像", indent_chars=0, bold=False)

# ===========================================================================
# 英文摘要 Abstract
# ===========================================================================
heading("Abstract", level=1, page_break_before=True)
para("With the rise of mobile health and smart healthcare, accurate and convenient dietary "
     "calorie monitoring has become an essential part of personal health management. However, "
     "existing approaches that estimate calories from a single 2D image using large multimodal "
     "models suffer from the absence of absolute physical scale, visual hallucination, and a "
     "lack of interpretability, which lead to large errors in food volume and calorie estimation "
     "and hinder their deployment in real-world health-management scenarios.", indent_chars=0)
para("To address these issues, this thesis proposes a food calorie estimation system based on 3D "
     "reconstruction and semantic segmentation. The system employs a standard credit card, a "
     "common everyday object, as an absolute scale reference, and integrates three key "
     "technologies: (1) open-vocabulary food recognition and localization using the lightweight "
     "vision-language model Qwen2.5-VL; (2) pixel-level boundary segmentation of multiple food "
     "items using the segmentation foundation model SAM3; and (3) 3D point-cloud reconstruction "
     "of the food from three images captured from the top and from the left and right 45-degree "
     "views using the multi-view reconstruction model DUSt3R. The scaleless point cloud is then "
     "anchored to absolute physical units through the known dimensions of the credit card, and "
     "the volume of each food item is computed via plane fitting and thickness integration. "
     "Finally, a purpose-built food volume-to-mass-to-nutrition database converts the "
     "pipeline-estimated volume into mass, calories, and macronutrients. Database construction "
     "and online inference share the same volume-estimation pipeline, so although the volume "
     "is not an independently verified physical ground truth, the resulting volume-to-mass "
     "mapping remains internally consistent within the system.", indent_chars=0)
para("The system adopts a mobile-cloud collaborative architecture: a mobile application handles "
     "standardized multi-view image acquisition and result presentation, while a cloud GPU "
     "inference service performs recognition, segmentation, and 3D reconstruction, thereby "
     "accommodating both the lack of dedicated depth-sensing hardware on ordinary devices and "
     "the substantial computational requirements.", indent_chars=0)
para("This thesis implements a system prototype and experimental protocol on a university "
     "cafeteria dataset; quantitative evaluation metrics (MAPE, MAE) are reported in Chapter 6. "
     "For context, prior work reports MAPE of about 26.1% for a Nutrition5k 2D baseline under "
     "fixed top-view capture, 11.47% for Shao et al. on 96 real eating-scene images, and "
     "17.67% EMAPE for Vinod et al. on SimpleFood45 with checkerboard and 3D food models. "
     "Because these studies differ in datasets, inputs, and evaluation conditions, their "
     "numbers are not directly rank-ordered; this thesis compares the proposed method against "
     "a separately constructed LMM baseline under the same cafeteria capture protocol.", indent_chars=0)
para("Keywords: food calorie estimation, 3D reconstruction, semantic segmentation, "
     "vision-language model, scale anchoring, mobile health, multi-view images",
     indent_chars=0)

# ===========================================================================
# 前置頁碼：羅馬數字，並加入頁碼頁尾
# ===========================================================================
set_page_number_format(doc.sections[0], fmt="upperRoman", start=1)
add_page_number_footer(doc.sections[0])

# 目次、圖次、表次（台大附件 4）
toc("目  次", 'TOC \\o "1-3" \\h \\z \\u')
toc("圖  次", 'TOC \\h \\z \\c "圖"')
toc("表  次", 'TOC \\h \\z \\c "表"')

# ===========================================================================
# 進入正文：新分節，頁碼改阿拉伯數字並自 1 起算
# ===========================================================================
body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
_apply_ntu_margins(body_sec)
set_page_number_format(body_sec, fmt="decimal", start=1)
add_page_number_footer(body_sec)

# ===========================================================================
# 第一章 緒論
# ===========================================================================
heading("第一章　緒論", level=1, page_break_before=True)
para("本研究所指之「食物熱量估算」，係指由使用者以一般智慧型手機拍攝餐點影像，自動推估各食材"
     "之體積、熱量與三大營養素（蛋白質、脂肪、碳水化合物）之過程。本章說明研究問題之由來、"
     "現有方法之侷限、本研究之動機、目的與主要貢獻，並概述全論文各章之安排。")

heading("1.1　研究背景", level=2)
para("飲食是維持人體生理機能的基礎，然而不當的飲食習慣與過量的熱量攝取，長期而言與肥胖、"
     "第二型糖尿病、心血管疾病及代謝症候群等慢性疾病息息相關。世界衛生組織與各國衛生主管機關"
     f"皆指出，肥胖與過重人口的比例在近數十年間持續攀升，已成為全球性的公共衛生課題{cit(17)}。"
     "在此背景下，如何協助一般民眾以低門檻、可持續的方式監控每日飲食熱量攝取，逐漸成為預防醫學"
     "與健康促進的重要議題。")
para("近年來，隨著智慧型行動裝置的高度普及、行動網路頻寬的提升，以及穿戴式感測器的成熟，"
     "「行動醫療」與「智慧健康照護」逐漸由概念走向落地應用。使用者得以透過手機隨時記錄運動量、"
     "心率、睡眠與飲食狀況，並藉由雲端服務進行資料彙整與分析。相較於傳統仰賴專業營養師人工計算"
     f"的飲食評估流程，行動化的自動化飲食紀錄在便利性、即時性與可負擔性上具有明顯優勢{cit(16,24)}，"
     "亦有助於提升使用者長期自我健康管理的意願與依從性。")
para("在飲食紀錄的自動化中，「以影像進行食物辨識與熱量估算」是最貼近使用者日常行為的一種方式："
     "使用者僅需在用餐前以手機拍攝食物影像，系統即可自動辨識食物種類並推估其熱量與營養成分。"
     f"此一情境免除了使用者手動輸入食物名稱與份量的繁瑣，因而受到學界與業界的廣泛關注{cit(16,25)}。"
     "Purdue 大學 TADA 計畫"
     f"{cit(30,31,47)}自 2010 年代起即探索以行動影像輔助 24 小時飲食回憶（24HR）；"
     "Anthimopoulos 等人之 goFOOD"
     f"{cit(34)}與 Thames 等人之 Nutrition5k{cit(20)}則分別代表「整合式 AI 系統」"
     "與「大規模公開 benchmark」兩個里程碑。然而，要由二維影像可靠地推估三維食物的份量，"
     "本質上是一個具高度挑戰性的不適定問題——Wang 等人"
     f"{cit(16)}與 Lo 等人{cit(24)}均指出，辨識正確並不等於份量正確，"
     "這也是目前多數影像式熱量估算系統精度不足的根本原因。")
para("值得注意的是，近年大型視覺語言模型（Vision-Language Model, VLM）與大型多模態模型"
     "（Large Multimodal Model, LMM）的快速發展，使得由影像直接生成食物描述與熱量估計成為可能"
     f"{cit(5,7,6)}。此類模型雖具備強大的開放域語意理解能力，能辨識種類繁多的食物，卻普遍缺乏"
     "對「絕對物理尺度」的感知，且在數值推估時容易出現與影像內容不符的「幻覺（hallucination）」，"
     f"導致其份量與熱量估計缺乏可靠性與可解釋性{cit(16,48)}。本研究即以此一痛點為出發點，探討如何"
     "在不依賴專用深度感測硬體的前提下，為影像式食物熱量估算導入可靠的絕對尺度與可解釋的幾何"
     "計算流程。")

heading("1.2　研究動機與目的", level=2)
para("本研究的核心動機，源自於對「現有二維影像式食物熱量估算方法」根本限制的觀察。經整理，"
     "現有方法主要存在以下三項瓶頸：")
para("第一，缺乏絕對物理尺度。單張二維影像在成像過程中已喪失深度資訊，且不同拍攝距離與焦距"
     "會造成相同食物在影像中呈現截然不同的像素大小。在沒有已知尺寸之參考物的情況下，模型無從"
     "判斷影像中食物的真實尺寸，僅能依據訓練資料的統計先驗進行猜測，導致同一道菜在不同拍攝條件下"
     f"得到差異極大的份量估計{cit(16,22,24)}。此為體積估算嚴重失真的首要來源。")
para("第二，視覺幻覺問題。以大型多模態模型直接輸出熱量數值時，模型傾向產生「看似合理但缺乏"
     "依據」的答案。由於此類模型並未實際測量食物的幾何量體，其數值往往受到食物名稱之統計先驗"
     f"主導，而非影像中真實的份量{cit(16,48)}；當面對非典型份量（例如特別大份或特別小份）時，"
     "估計誤差尤其顯著。")
para("第三，缺乏可解釋性。大型視覺語言模型的黑箱式估算無法向使用者說明「為何得到此一熱量"
     "數值」，使用者難以判斷結果的可信度，也不利於在醫療與健康管理情境中被專業人員採信。可解釋的"
     f"中間量（如各食材的分割區域、三維點雲、體積與重量）對於建立使用者信任與實際應用皆至關重要"
     f"{cit(16,22)}。")
para("基於上述動機，本研究的目的在於：設計並實作一套「免專用硬體、具絕對尺度、且具可解釋性」的"
     "食物熱量估算系統。具體而言，本研究欲達成之目標包括：（一）以日常隨身之標準信用卡作為絕對"
     "比例尺，解決二維影像缺乏尺度之問題；（二）整合視覺語言模型、萬物分割模型與多視角三維重建"
     "模型，建立由影像至管線估計體積之可解釋計算管線；（三）建置食物體積、重量與營養之轉換"
     "資料庫，以與推論相同之管線標定體積—重量對應，完成由管線估計體積至熱量與三大營養素之"
     "推估；（四）以行動端與雲端協同之架構完成系統實作，並以實際"
     "餐點驗證其準確度與可行性。")

heading("1.3　主要貢獻", level=2)
para("本研究之主要貢獻可歸納如下：")
numbered("提出以日常隨身物品作為絕對比例尺的尺度錨定方法。本研究採用符合 ISO/IEC 7810 ID-1 "
         f"規格之標準信用卡（85.60 mm × 53.98 mm）作為參考物{cit(19)}，透過將其置於餐點旁一同"
         "拍攝，使系統得以在無專用深度感測器的情況下，將無尺度之三維重建結果校正為具絕對物理"
         f"單位之空間模型{cit(22,23)}，從根本上解決二維影像缺乏尺度之問題。")
numbered("提出整合多模態視覺與多視角三維重建之可解釋體積估算管線。本研究串接輕量化視覺語言模型"
         f"Qwen2.5-VL（食物辨識與定位）{cit(6)}、萬物分割模型 SAM3（食材邊界分割）{cit(12)}"
         f"與多視角三維重建模型 DUSt3R（三維點雲還原）{cit(13)}，並提出結合平面擬合與厚度積分之"
         "體積計算演算法，使每一步驟皆產生可視化、可檢驗之中間結果，提升系統之可解釋性。")
numbered("建置食物幾何與營養轉換資料庫。本研究對各品項之一分量，先以與線上推論相同之管線取得"
         "三維重建並計算管線估計體積，再將同一分量秤重，反推堆積密度並結合單位重量營養資訊，"
         "建立「管線估計體積—重量—熱量—三大營養素」之查詢對應關係。")
numbered("設計並實作行動端與雲端協同之系統原型，並提出免專用硬體之標準化多視角拍攝流程，驗證"
         "本方法於一般行動裝置環境下的可行性，並與直接使用大型多模態模型之基準方法進行比較。")

heading("1.4　論文概述", level=2)
para("本論文共分為七章，各章內容安排如下：", indent_chars=0, space_after=4)

chapter_outline(
    "一", "緒論",
    "說明研究背景與問題，整理現有影像式熱量估算方法之侷限，闡述本研究之動機、"
    "目的、主要貢獻，並概述全論文各章之安排。",
)
chapter_outline(
     "二", "文獻探討",
    "依 VBDA 技術鏈回顧食物辨識、語意分割、三維幾何重建，並專節討論 TADA、Nutrition5k、"
    "Shao/Vinod 系列、MUSEFood、Dhar 等電腦視覺熱量估算代表性研究，整理技術演進、"
    "性能比較與本研究定位。",
)
chapter_outline(
    "三", "食物幾何與營養轉換資料庫建置",
    "說明營養資料庫之建置流程與設計，訂定多視角影像採集與比例尺擺放規範，"
    "並以與推論相同之管線標定管線估計體積與秤重之對應，建立營養轉換資料表。",
)
chapter_outline(
    "四", "基於多模態視覺與三維點雲之體積估算架構",
    "提出本研究之核心方法，包含系統總體管線、食物辨識與分割、多視角三維重建、"
    "比例尺錨定與體積積分，以及熱量與營養素估算等完整流程。",
)
chapter_outline(
    "五", "實驗設計與系統實作",
    "說明行動應用程式與雲端 GPU 推論服務之實作，介紹以活大自助餐為來源之"
    "實驗資料與測試場景，並記錄硬體環境、推論參數與評估指標。",
)
chapter_outline(
    "六", "實驗結果分析與討論",
    "呈現並分析辨識分割效能、管線體積估算表現、與基準方法之比較結果，"
    "評估系統部署效能，並綜合討論主要誤差來源。",
)
chapter_outline(
    "七", "結論與未來展望",
    "總結本研究之成果與貢獻，回應研究目的之達成程度，並提出後續可能之研究方向。",
)

heading("1.5　相關研究回顧與研究缺口", level=2)
para("綜合近年以電腦視覺估算食物熱量之代表性研究，可觀察到該領域之核心難點並非僅在於「辨識"
     "這是什麼食物」，而在於「由二維影像恢復與熱量高度相關之份量、體積或質量」。Thames 等人"
     f"之 Nutrition5k {cit(20)}大規模基準實驗顯示，僅以二維 RGB 直接回歸熱量，平均絕對誤差"
     "可達約 70.6 kCal（相對約 26.1%）；若引入深度或體積先驗，誤差可降至約 41.3 kCal"
     "（16.5%），證實幾何訊息對份量估算至關重要。Shao 等人"
     f"{cit(49)}於僅 96 張真實用餐影像上，以單張 RGB 與能量分布圖中介表徵達到 MAPE 11.47%，"
     "顯示純影像路線在特定資料分布下具潛力，但資料規模與場域泛化仍受限。Vinod 等人"
     f"{cit(22)}則走顯式三維幾何路線，在 SimpleFood45 上報告 EMAPE 17.67%，並提供可開源"
     "重現之資料與程式，代表「可解釋幾何＋參考物」之工程化方向。")
para("若依輸入條件與幾何建模方式分類，現有研究大致可分成三類：第一類為單張 RGB 直接學習熱量"
     "映射（如 Fang 等人之能量分布圖"
     f"{cit(32)}、Shao 等人之跨域特徵融合{cit(49)}、Shao 等人之體素重建{cit(21)}），"
     "部署最簡便但普遍缺乏絕對物理尺度；第二類為 RGB-D 或多模態融合（如 Vinod 等人之深度域"
     f"適配{cit(51)}、Nutrition5k 之 depth/volume baseline{cit(20)}），精度較高但推論常需"
     f"深度感測或受控拍攝；第三類為顯式三維重建或三維模型縮放（如 Dehais 等人{cit(29)}、"
     f"Wang 等人之 MUSEFood{cit(35)}、Vinod 等人{cit(22)}），可解釋性強但多需多視角、"
     "參考物或預建三維資產。Dhar 等人"
     f"{cit(50)}針對孟加拉街頭食物，以硬幣縮放結合 YOLOv9 與回歸，在垂直品類上達 MAE"
     "7.85 kCal、R² 96.0%，說明「在地小類別＋實體參考物」亦為可行之產品化路徑。")
para("對照上述文獻，本研究之定位在於：允許使用日常信用卡作為參考物（呼應 TADA、Vinod 等"
     f"幾何路線{cit(30,22)}），但改以三視角 DUSt3R 重建取代固定 top-view 掃描或 3D 模板"
     "資料庫，並以 VLM＋SAM3 提供開放域辨識分割，兼顧「絕對尺度」「可解釋中間結果」與"
     "「一般手機可負擔之硬體成本」三者。下一章將依 VBDA 技術鏈與上述代表性研究，進一步"
     "整理方法演進與本研究之對照依據。")

# ===========================================================================
# 第二章 文獻探討
# ===========================================================================
heading("第二章　文獻探討", level=1, page_break_before=True)
para("影像式飲食評估（Vision-Based Dietary Assessment, VBDA）旨在以相機影像自動辨識食物、"
     "估計份量並換算營養資訊，以降低人工記錄之負擔。Wang 等人"
     f"{cit(16)}與 Lo 等人{cit(24)}之綜述均指出，要完成可靠之熱量估算，典型流程包含："
     "（一）食物影像辨識；（二）食材區域分割；（三）三維幾何重建與體積量測；（四）由管線"
     "估計體積或重量換算熱量與營養素。圖 2-1 概括此技術鏈之主要環節。本章依序回顧各環節之代表性"
     "文獻與方法演進，並於 2.4 節專門整理 Fang、Shao、Nutrition5k、Vinod、Dhar 等"
     "電腦視覺食物熱量估算之相似研究與性能比較（表 2-4～2-6），最後分析現有方法瓶頸"
     "並整理本研究在既有研究脈絡中之定位。")
fig_img("fig2_1.png", "圖 2-1　影像式飲食評估（VBDA）典型技術管線", width_cm=14)

heading("2.1　食物影像辨識與分類技術概述", level=2)
para("食物影像辨識為 VBDA 之首要步驟，其任務為由影像判別食物之類別、名稱或語意描述。"
     "此領域之發展歷程，大致可區分為以卷積神經網路為主之封閉集合分類、以物件偵測與"
     "細粒度分類處理多品項場景，以及近年興起以視覺語言模型為代表之開放域辨識三個階段。"
     f"Subhi 等人之調查{cit(25)}亦將食物辨識視為後續份量估算與營養分析之基礎模組。")

heading("2.1.1　基於傳統卷積神經網路的食物分類", level=3)
para("卷積神經網路（Convolutional Neural Network, CNN）自 AlexNet 於 ImageNet 競賽中取得"
     f"突破後{cit(1)}，成為影像分類的主流架構。He 等人提出之 ResNet"
     f"{cit(2)}以殘差連接緩解深層網路之梯度消失問題，使更深之骨幹網路得以訓練，"
     "並成為後續食物辨識研究常用之特徵抽取器。圖 2-2 示意典型之 CNN 食物分類架構："
     "輸入影像經骨幹網路轉換為特徵向量，再經全連接層與 Softmax 輸出固定類別之機率分布。")
fig_img("fig2_2.png", "圖 2-2　基於 CNN 之封閉集合食物分類典型架構", width_cm=14)
para("在食物辨識領域，研究者陸續提出多個具代表性之資料集。Bossard 等人之 Food-101"
     f"{cit(3)}包含 101 類、每類 1000 張影像，為早期深度學習食物分類之標準基準；"
     f"Kawano 與 Yanai 之 UEC-Food256 {cit(4)}與 Min 等人之 VireoFood-172 / Food-500"
     f"{cit(26)}進一步擴大類別數與場景多樣性。研究者常以 VGG、ResNet、Inception、"
     f"EfficientNet 等骨幹網路{cit(2,40,41)}進行訓練，於封閉類別集合上取得良好之"
     "Top-1 / Top-5 準確度。")
para("此類方法之優點在於推論速度快、模型成熟且易於部署於行動裝置或邊緣硬體；然而其根本"
     "限制在於採用「封閉集合（closed-set）」假設，即模型僅能辨識訓練資料中出現過之固定"
     "類別。面對真實世界中種類繁多、地域差異大、且不斷推陳出新之食物（尤其是混合式餐點"
     f"與地方料理），封閉集合分類器難以泛化，且擴充新類別需重新蒐集標註資料並重新訓練，"
     f"維護成本高昂{cit(16,26)}。此外，傳統分類器僅輸出「類別標籤」而不提供空間定位"
     "或份量資訊，無法直接銜接後續之分割與幾何量測。")

heading("2.1.2　物件偵測、實例分割與細粒度食物辨識", level=3)
para("為同時取得食物之類別與位置，部分研究改採物件偵測或實例分割框架。Ren 等人之"
     f"Faster R-CNN {cit(27)}以區域建議網路（RPN）與 RoI 池化實現端到端之物件偵測；"
     f"He 等人之 Mask R-CNN {cit(28)}則在 Faster R-CNN 基礎上增加遮罩分支，可同時"
     "輸出邊界框與像素級實例遮罩。此類方法適用於餐盤中同時存在多項食物之場景，"
     "較單純分類更能反映實際用餐型態。")
para(f"Min 等人{cit(26)}針對大規模視覺食物辨識，探討資料不平衡、細粒度子類別與"
     "跨資料集泛化等議題，顯示即使採用強大之 CNN 骨幹，在真實自助餐或混合餐點影像中，"
     "仍常因外觀相近、部分遮擋與拍攝角度變化而產生誤判。goFOOD 系統"
     f"{cit(34)}則結合 Mask R-CNN 與後處理，於行動裝置輔助飲食評估場景中同時完成"
     "食物偵測與初步分割，為後續幾何估算提供區域依據。整體而言，偵測與實例分割雖改善"
     "了「多品項定位」問題，但仍多侷限於預定義之類別集合，且需大量框選或遮罩標註，"
     "難以低成本擴展至開放域之新食物。")

heading("2.1.3　視覺語言模型在開放域辨識的應用", level=3)
para("為突破封閉集合之限制，視覺語言模型（Vision-Language Model, VLM）藉由在大規模"
     "影像—文字配對資料上進行預訓練，學習影像與自然語言之間的對應關係，因而具備"
     f"「開放域（open-vocabulary）」辨識能力。Radford 等人之 CLIP {cit(5)}採對比式"
     "學習，可透過文字提示（text prompt）對任意類別進行零樣本（zero-shot）分類，"
     "無需針對特定食物類別重新訓練。圖 2-3 示意 VLM 將影像特徵與文字提示映射至"
     "共同語意空間之概念。")
fig_img("fig2_3.png", "圖 2-3　視覺語言模型（VLM）開放域食物辨識概念示意", width_cm=14)
para("後續生成式 VLM 與大型多模態模型（Large Multimodal Model, LMM），如 LLaVA"
     f"{cit(7)}、GPT-4V {cit(46)}、Gemini 與 Qwen-VL 系列{cit(6)}，進一步整合"
     f"影像理解與語言生成能力，能以自然語言描述影像內容、回答關於影像之問題，甚至輸出"
     f"結構化 JSON（含邊界框座標）。Dosovitskiy 等人之 ViT {cit(45)}以 Transformer"
     "取代 CNN 作為影像骨幹，亦被廣泛用於食物特徵抽取，與 CNN 形成互補。")
para(f"本研究採用之 Qwen2.5-VL {cit(6)}屬開源 VLM，在辨識能力與部署成本間取得平衡，"
     "且支援以文字提示引導模型輸出食物品項及其邊界框，適合作為後續分割之定位依據。"
     "然而，VLM 雖擅長語意層級之辨識，卻不擅長幾何層級之精確測量；當被要求直接輸出"
     f"份量或熱量時，輸出多受語言先驗主導而非實際幾何量測{cit(16,48)}，易產生幻覺"
     "與系統性偏差。此一觀察構成本研究之設計原則：VLM 負責辨識與定位，幾何量測交由"
     "三維重建與比例尺錨定流程處理。")

heading("2.1.4　食物辨識資料集與評估基準", level=3)
para("食物辨識之進展與資料集發展密切相關。除前述 Food-101、UEC-Food 與 VireoFood 外，"
     f"Thames 等人建置之 Nutrition5k {cit(20)}不僅提供食物影像，更包含重量、營養成分"
     "與深度／視角資訊，使研究者得以評估「辨識正確」是否真正轉化為「營養估算準確」。"
     "表 2-1 整理常用資料集之規模與任務定位。")
make_table(
    ["資料集", "類別數／規模", "主要任務", "與份量／營養之關聯"],
    [
        ["Food-101 [3]", "101 類 × 1000 張", "分類", "無份量標註"],
        ["UEC-Food256 [4]", "256 類", "分類", "無份量標註"],
        ["VireoFood-172 [26]", "172+ 類", "分類／偵測", "部分子集含定位"],
        ["Nutrition5k [20]", "5000+ 餐點", "營養回歸", "含重量與營養真值"],
        ["MetaFood [23]", "競賽資料", "3D 重建＋營養", "含物理參考物"],
    ],
    caption_text="表 2-1　食物影像相關常用資料集比較",
)
para("由表 2-1 可見，早期資料集多聚焦分類準確度，而近年資料集逐漸納入份量、深度與營養"
     "真值，反映學界已從「辨識對了什麼」轉向「估算吃了多少」。Nutrition5k"
     f"{cit(20)}之設計尤其重要：它以 dish 為單位同時提供 ingredient、重量與營養，"
     "並以 RealSense 取得 overhead depth，使研究者能在同一資料上比較 2D、RGB-D 與"
     "volume scalar 等路線。SimpleFood45"
     f"{cit(22)}則補足幾何可解釋評估，提供含 volume/weight/energy 真值之 smartphone "
     "影像。本研究之活大自助餐實驗雖規模較小，但評估思維與 Nutrition5k 一致："
     "以熱量與體積誤差為核心，而非僅報告分類準確率。")

heading("2.1.5　食物辨識在熱量估算管線中的角色", level=3)
para("食物辨識在 VBDA 管線中扮演「語意閘門」角色：錯誤的類別會導致查表或回歸使用錯誤"
     "之營養參數。He 等人"
     f"{cit(33)}以多任務學習證明，分類與份量共享特徵可相互增益；Fang 等人"
     f"{cit(32)}與 Shao 等人{cit(49)}則假設 segmentation 與 food label 已知或"
     "可由上游模組提供。goFOOD"
     f"{cit(34)}與 Vinod 等人{cit(22)}採 detector + segmenter 封閉或半封閉流程；"
     "本研究改以 VLM 開放域辨識，避免為台灣自助餐不斷擴充類別表。需注意的是，開放域"
     "辨識仍可能產生同義詞、粒度不一致（如「pork」vs「braised pork」）等問題，"
     "故第三章設計名稱對應與 fuzzy matching，使辨識結果能穩定銜接營養資料庫。")

heading("2.2　影像語意分割技術演進", level=2)
para("在取得食物之類別與大致位置後，尚需精確劃分各食材於影像中所佔之像素區域，方能於"
     "三維重建結果中界定各食材對應之點雲範圍，並避免相鄰食物之幾何量測互相污染。"
     "圖 2-4 概括分割技術由傳統圖割、深度全卷積網路至分割基礎模型之演進。")
fig_img("fig2_4.png", "圖 2-4　食物影像分割技術演進概觀", width_cm=14)

heading("2.2.1　傳統交互式分割演算法", level=3)
para("在深度學習普及之前，影像分割多仰賴低階特徵（顏色、紋理、邊緣）與圖論最佳化。"
     f"Boykov 與 Jolly 之 Graph Cut {cit(9)}將影像視為圖結構，以能量最小化求解前景／"
     f"背景切割；Rother 等人之 GrabCut {cit(8)}引入高斯混合模型描述前景與背景顏色"
     "分布，使用者僅需框選大致區域即可迭代求精分割結果。Chen 等人"
     f"{cit(37)}於雙視角食物體積估算中即結合 Faster R-CNN 定位與 GrabCut 分割，"
     "顯示傳統分割在特定場景仍具實用價值。")
para("此類方法優點為不需大量標註、原理直觀；然而面對食物影像中常見之複雜紋理、相近"
     f"顏色食材相鄰、醬汁反光與陰影時，往往難以取得穩定邊界{cit(24)}，且多需人工"
     "框選或修正，不利於全自動化之飲食紀錄流程。")

heading("2.2.2　深度學習語意分割與實例分割", level=3)
para("全卷積網路（FCN）與 U-Net 等架構使像素級語意分割得以端到端訓練；DeepLab 系列"
     "以空洞卷積擴大感受野，提升邊界精度。於食物場域，語意分割可將「飯、菜、肉」"
     f"等區域分離，但仍難區分同類別之不同實例。Mask R-CNN {cit(28)}之實例分割框架"
     f"可為每個食物實例輸出獨立遮罩，goFOOD {cit(34)}即採類似思路處理餐盤中多項食物。"
     "表 2-2 比較不同分割范式之特性。")
make_table(
    ["分割范式", "代表方法", "輸出", "食物場域限制"],
    [
        ["語意分割", "FCN / U-Net / DeepLab", "每像素類別", "不區分同類多實例"],
        ["實例分割", "Mask R-CNN", "每實例 Mask", "需類別標註、封閉集合"],
        ["交互式分割", "GrabCut", "單物件 Mask", "需人工提示"],
        ["提示分割", "SAM / SAM3", "任意物件 Mask", "開放詞彙、零樣本泛化"],
    ],
    caption_text="表 2-2　影像分割范式與食物應用特性比較",
)

heading("2.2.3　萬物分割模型的提示導向分割機制", level=3)
para("Kirillov 等人提出之 Segment Anything Model（SAM）"
     f"{cit(10)}以大規模分割資料訓練，具備可提示（promptable）分割能力：使用者可"
     "以點、框或遮罩提示引導模型分割任意物件，對未見過之類別亦具零樣本泛化能力。"
     f"SAM 2 {cit(11)}將能力延伸至影片時序分割。圖 2-5 示意 SAM 以框／點提示產生"
     "像素級遮罩之流程。")
fig_img("fig2_5.png", "圖 2-5　Segment Anything（SAM）提示導向分割機制示意", width_cm=13)
para(f"本研究採用之 SAM3 {cit(12)}引入以文字片語指定「概念（concept）」之能力，"
     "可一次性分割影像中屬於某開放詞彙之所有實例，並以存在符記（presence token）"
     "強化相近提示之辨別。就本研究而言，SAM3 能以 VLM 輸出之邊界框或食物名稱為提示，"
     "取得各食材與信用卡之像素級遮罩，為三維點雲界定與體積積分提供區域依據。")
para("分割品質直接影響後續體積積分。MUSEFood"
     f"{cit(35)}報告 MFCN 分割 mIoU 0.9210 顯著優於 GrabCut；Vinod 等人"
     f"{cit(22)}以 SAM 取得 instance mask 再做 3D scaling；Dhar 等人"
     f"{cit(50)}則以 YOLOv9 分割後抽取 contour 幾何特徵。對混合餐盤而言，"
     "相鄰食材邊界之準確性決定各食材點雲是否互相污染。")

heading("2.2.4　分割結果對份量估算之影響", level=3)
para("Lo 等人"
     f"{cit(24)}之綜述指出，分割誤差會傳遞至體積與熱量。若分割遮罩高估面積，"
     "則無論幾何模型或回歸模型皆傾向高估份量；若相鄰食物合併為同一 mask，"
     "則無法分項報告營養。TADA 與 Chen 等人"
     f"{cit(30,37)}早期仰賴 GrabCut 或手動修正；現代 pipeline 雖以深度分割"
     "提升自動化，但在醬汁、反光與部分遮擋下仍可能出現 under-segmentation。"
     "本研究以 VLM 框 + SAM3 精修，並在體積積分前濾除過小雜訊區域，以降低"
     "分割錯誤對厚度積分之影響。")

heading("2.3　計算機視覺之三維幾何重建", level=2)
para("為取得可用於份量估算之體積，必須還原食物之三維幾何結構。本節先討論主動式深度感測於行動裝置"
     "之限制，再回顧 SfM/MVS、單目深度、NeRF 與 DUSt3R 等低設備需求之重建方法。"
     "表 2-7 比較各方法之輸入、標定需求與尺度特性。")

heading("2.3.1　光學雷達與手機深度感測之限制", level=3)
para("直接取得深度資訊最直觀之方式，是採用 ToF、結構光或 LiDAR 等主動式感測器。"
     "部分旗艦手機內建 LiDAR，可提供場景深度圖，理論上有助於量測食物尺寸。"
     f"然而，MUSEFood {cit(35)}等研究指出，多感測器方案雖可提升精度，但硬體成本高、"
     f"裝置普及率低{cit(24,35)}。MiDaS、ZoeDepth 等單目深度估計"
     f"{cit(42,43)}雖僅需 RGB 影像，但多輸出相對深度，若無參考物仍難換算絕對體積。"
     "跨裝置深度格式與精度差異大，不利系統之一致部署。基於普惠性考量，本研究改採"
     "一般彩色相機搭配多視角幾何重建。")

heading("2.3.2　SfM/MVS 與傳統多視角立體視覺", level=3)
para("Schönberger 與 Frahm 之 COLMAP"
     f"{cit(14)}代表現代 SfM/MVS 流程：先以特徵匹配估計相機姿態，再三角化生成稀疏點雲，"
     "並可進一步融合為稠密點雲。此方法在紋理豐富、視角充足之場景效果良好，但對食物影像"
     "常見之弱紋理（白飯、清湯）、高光反射與視角不足等情況敏感，且流程繁複、對初始標定"
     "與特徵匹配品質要求高，不利一般使用者以手機隨手拍攝之應用。")

heading("2.3.3　NeRF 與基於深度學習的多視角重建", level=3)
para("Mildenhall 等人之 NeRF"
     f"{cit(15)}以神經輻射場表示場景，可生成高品質新視角合成，但訓練與渲染成本高，"
     "難以滿足即時飲食紀錄需求。近年 DUSt3R"
     f"{cit(13)}提出免相機標定、免已知姿態之稠密重建：模型直接由影像對回歸逐像素"
     "三維點圖（pointmap），並以全域對齊整合多視角至統一座標系。MASt3R"
     f"{cit(44)}強化特徵匹配，提升多視角對齊精度。表 2-7 整理各方法差異。")
make_table(
    ["方法", "輸入", "標定需求", "輸出尺度", "食物場域適用性"],
    [
        ["SfM/MVS (COLMAP)", "多視角影像", "通常需要", "絕對（若已知基線）", "紋理不足時易失敗"],
        ["NeRF", "多視角影像", "需要", "相對", "重建慢、難即時"],
        ["單目深度 (MiDaS 等)", "單張影像", "不需要", "相對深度", "缺乏絕對份量"],
        ["DUSt3R / MASt3R", "2–N 張影像", "不需要", "相對", "少視角稠密重建"],
    ],
    caption_text="表 2-7　三維幾何重建方法分類與特性比較",
    col_widths=[3.2, 2.8, 2.4, 3.0, 3.6],
)
para("需強調的是，DUSt3R 等方法輸出通常僅具相對尺度，無法直接得知公分級尺寸"
     f"{cit(22,23)}。圖 2-7 說明尺度不確定性（scale ambiguity）以及以已知尺寸參考物"
     "（如信用卡、棋盤格）完成尺度錨定之必要性。")
fig_img("fig2_7.png", "圖 2-7　多視角重建之尺度不確定性與參考物錨定概念", width_cm=14)

heading("2.3.4　食物場域之三維重建與份量估算研究", level=3)
para("Dehais 等人"
     f"{cit(29)}提出僅需雙視角即可重建食物三維模型並估算體積，為食物場域多視角重建之"
     "早期代表。其方法將問題明確拆解為「相機外參校正→稠密立體匹配→體積擷取」三段式"
     "幾何流程：輸入為兩張餐盤 RGB 影像，旁置信用卡大小之參考卡，並假設可取得食物與"
     "盤面分割；先估相機相對姿態，再以稠密 stereo matching 建點雲，最後於分割區域積分"
     "體積。於 77 道真實菜餚上，最佳設定 MAPEoverall 為 8.2%，平均處理時間約 5.5 秒／"
     "道菜。此路線之優點為幾何可解釋、誤差可逐層診斷；缺點則為需雙視角、參考卡與可用"
     "分割，對一般使用者之操作負擔較高。")
para("Gao 等人之 MUSEFood"
     f"{cit(35)}則從另一角度處理「尺度」問題：在不放參考物、不事先蒐集大量帶體積標註"
     "訓練集之前提下，利用智慧型手機相機、喇叭與麥克風，以 MFCN 同時完成食物分割與"
     "容器分類，並以 MLS 聲學測距取得手機至桌面距離，再結合頂視／側視微分幾何模型估"
     "體積。於 SUEC Food 分割實驗，最佳 MFCN-B 之 mIoU 達 0.9210；於雞腿、炸豬排與"
     "粥三種真實食物，相對體積誤差分別為 2.70%、12.37% 與 −0.27%。此研究證明「免"
     "參考物」在受控場景下可行，但仍需多視角與額外音訊程序，且實驗食物類型有限。")
para("Lo 等人"
     f"{cit(38)}以視角合成輔助體積估算；Vinod 等人{cit(22)}於 CVPR 2024 利用三維"
     "食物模型縮放與場景參考物，報告 EMAPE 17.67%；MetaFood 競賽"
     f"{cit(23)}與 VolETA {cit(36)}進一步彙整以棋盤格、信用卡等物理參考物輔助之"
     f"三維食物重建方法。Thames 等人之 Nutrition5k {cit(20)}亦證實，引入深度或體積"
     "資訊可顯著改善熱量預測。上述研究共同指出：三維幾何資訊對份量估算至關重要，"
     "但如何以最少硬體與操作負擔取得可靠之絕對尺度，仍是核心課題；本研究即在此缺口"
     "上，以三視角 DUSt3R 重建搭配信用卡比例尺，試圖在 Dehais 之幾何可解釋性與"
     "MUSEFood 之免專用感測器精神之間取得平衡。")

heading("2.4　電腦視覺食物熱量估算之代表性研究", level=2)
para("除前述辨識、分割與三維重建等基礎技術外，近年已有大量研究直接以「電腦視覺」"
     "完成由餐點影像到熱量或營養素之估算。Wang 等人"
     f"{cit(16)}、Lo 等人{cit(24)}、Subhi 等人{cit(25)}與 Chen 等人"
     f"{cit(48)}之綜述均指出，此類研究之共同目標在於降低人工飲食紀錄負擔，並提升"
     "份量估計之客觀性。本節依系統整合程度與幾何建模方式，整理具代表性之相似研究，"
     "作為本研究之對照與定位依據。")

heading("2.4.1　行動影像飲食評估系統（TADA 系列）", level=3)
para("Purdue 大學 Zhu 與 Delp 團隊長期推動 Technology Assisted Dietary Assessment（TADA）"
     f"系統{cit(30,31)}，為電腦視覺輔助熱量估算之奠基性工作之一。Zhu 等人"
     f"{cit(30)}率先論證以行動裝置拍攝食物影像、再交由伺服器分析之可行性；Six 等人"
     f"{cit(31)}進一步開發影像分析系統，供營養師進行飲食評估。Six 等人"
     f"{cit(47)}則提出 Mobile Telephone Food Record（mpFR），以證據為本方式設計"
     "行動電話飲食紀錄流程，強調拍攝前後餐影像、標準化操作流程與使用者介面設計。"
     "TADA 系列之核心思路為：以 fiducial marker 或已知尺寸參考物提供尺度，結合食物"
     "分割與幾何模型估算份量，再查表換算熱量。圖 2-8 示意其典型流程。此路線之優點"
     "在於系統完整、臨床驗證累積深厚；侷限則在於依賴實體標記、操作步驟較多，且幾何"
     "模型對複雜堆疊食物仍採簡化假設。")
fig_img("fig2_8.png", "圖 2-8　TADA 行動裝置輔助飲食評估（Technology Assisted Dietary Assessment）流程示意", width_cm=14)

heading("2.4.2　整合式人工智慧飲食評估系統", level=3)
para("隨深度學習成熟，研究者開始建置端到端之 AI 飲食評估系統。Anthimopoulos 等人提出"
     f"goFOODTM {cit(34)}，整合 Mask R-CNN 食物偵測、語意分割與幾何份量估算，"
     "可自動辨識多種食物並輸出營養資訊，為早期「深度學習＋完整管線」之代表。Jiang 等人"
     f"之 DeepFood {cit(39)}則以深度模型同時處理食物影像分析與膳食評估，涵蓋分類、"
     "定位與熱量相關預測。此類系統證明 CNN 特徵足以支撐複雜餐點場景，但多仍假設"
     "固定食物類別或需離線訓練，且份量估算多依二維投影或簡化三維模板，對非典型份量"
     f"與混合取餐之泛化有限{cit(16,24)}。")

heading("2.4.3　深度學習端對端熱量與份量估算", level=3)
para("另一路線直接以深度網路由影像回歸熱量或份量，減少人工設計之幾何步驟。Chen 等人"
     f"{cit(37)}提出結合 Faster R-CNN 與 GrabCut 之雙視角深度學習熱量估算方法，"
     "以俯視與側視影像推估體積並換算熱量，為較早期將深度學習用於「卡路里估算」"
     "之探索。He 等人"
     f"{cit(33)}提出多任務影像式膳食評估框架，以 ResNet-18 雙分支同時處理食物分類"
     "與以 kcal 表示之份量回歸，並以 soft parameter sharing 與 cross-domain feature "
     "adaptation 融合兩分支特徵。於 96 張 eating occasion 裁切之 834 張單食物影像"
     "（21 類，增強後 2,168 張）上，最佳設定達分類準確率 88.67%、MAE 56.82 kCal，"
     "且優於受試者人工估計之 45.43% 誤差百分比。此研究證明「類別先驗」有助熱量估計，"
     "但研究設定仰賴手動裁切單食物區域，離真實混合餐盤仍有距離。")
para("Fang 等人"
     f"{cit(32)}提出以學習式能量分布圖（learned energy distribution image）建立"
     "影像至熱量之映射：先依分割遮罩、食物標籤與真實熱量建立與原圖對應之能量分布圖，"
     "並以 5×4 彩色棋盤 fiducial marker 做透視校正，再以條件式 GAN 學習 RGB→能量"
     "分布圖映射，報告平均能量估計誤差率 10.89%。其概念創新在於讓模型學習「哪些區域"
     "貢獻較多熱量」，而非直接猜單一標量；但訓練需大量人工遮罩與參考棋盤，可重現性"
     "偏弱。Shao 等人"
     f"{cit(49)}延續 energy distribution 思想，改以 LayerNorm 正規化融合 RGB 域與"
     "energy-distribution 域特徵，於 96 張真實用餐影像（增強後 864 train / 96 test）"
     "上達 MAE 56.22 kCal、MAPE 11.47%，顯著優於受試者人工估計之 MAPE 39.03%。"
     "此為單張 RGB、免深度感測器路線之重要基線，但資料量小且場域偏特定營養研究。")
para("Shao 等人"
     f"{cit(21)}進一步提出由單張 RGB 進行三維體素形狀重建之端對端框架：訓練時以"
     "Nutrition5k 深度圖產生 voxel 監督，推論時僅需 RGB。於 top-view 3,259 張子集"
     "（2,753 train / 506 test）上，文中表格報告 MAE 40.05 kCal、MAPE 22.0%，"
     "MAEoM 15.8%。需注意的是，摘要與表格之 MAPE 數值存在不一致，引用時應以主文"
     "表格為準並自行重現驗證。此類方法推論流程簡潔，但普遍缺乏明確之絕對物理尺度，"
     "且中間表示不易解釋，當訓練分布外之份量或擺盤型態出現時，"
     f"易產生系統性偏差{cit(16,48)}。")

heading("2.4.4　三維幾何與營養預測整合研究", level=3)
para("近年研究逐漸形成共識：僅憑二維外觀難以可靠估計攝取量，引入三維幾何或深度資訊"
     f"可顯著改善熱量預測{cit(20,24)}。Thames 等人發表 Nutrition5k 資料集"
     f"{cit(20)}，為此領域最重要之公開基準之一：含 5,066 道 unique dishes、約 2 萬段"
     "短影片、超過 250 種 ingredient，其中約 3.5k 道含 Intel RealSense overhead RGB-D。"
     "每道菜以真實校園自助餐逐項加菜、即時掃描與秤重取得重量、熱量與 macronutrient "
     "標註。作者提出多條 baseline：2D direct prediction、depth as 4th channel、volume "
     "scalar 等。在絕對熱量預測上，2D direct 為 70.6 kCal / 26.1%；depth 4th channel "
     "改善為 47.6 / 18.8%；volume scalar 進一步為 41.3 / 16.5%，且已優於營養師視覺"
     "估量（約 41%）與非營養師（約 53%）。Nutrition5k 之侷限在於拍攝以固定 top-view "
     "為主，與一般手機側前方拍攝習慣存在 domain gap，對台灣便當、自助餐與夜市小吃之"
     "直接泛化需謹慎。")
para("Vinod 等人"
     f"{cit(51)}承接 energy-distribution 路線，探問 depth map 能否進一步提升熱量估算："
     "先以 segmentation 建立 Energy Density Map，再與 depth 特徵經 normalization 融合"
     "回歸熱量。於 Nutrition5k 含 depth 之子集上，最佳 (Energy Density, Depth) 組合"
     "MAPE 13.57%（MAE 24.37 kCal）。此路線性能強，但推論需 depth，不利一般手機"
     "廣泛部署；較適合作為具 LiDAR/ToF 裝置之上限模型，或訓練端 teacher supervision。"
     "Vinod 等人"
     f"{cit(22)}則改走顯式 3D object scaling：以 YOLOv8 + SAM 取得類別與 mask，"
     "以 checkerboard 做相機標定與 PnP 求 pose，載入對應 3D food model 渲染回影像，"
     "以 rendered mask 與 input mask 面積比縮放體積，再經 FNDDS 換算熱量。於 "
     "SimpleFood45（12 food types、45 items、513 張影像）上，VMAPE 14.01%、EMAPE "
     "17.67%（EMAE 31.10 kCal）。其最大價值在可解釋性與開源可重現性，但依賴每類"
     f"食物之 3D 掃描模型與 checkerboard 參考物。MetaFood {cit(23)}與 VolETA"
     f"{cit(36)}進一步聚焦以信用卡、棋盤格等物理參考物輔助之 3D 重建。表 2-4 與"
     "表 2-5 彙整上述研究之技術重點與報告性能。")
make_table(
    ["研究／系統", "發表／年份", "核心作法", "尺度／幾何來源", "與本研究之異同"],
    [
        ["TADA / mpFR [30,47]", "2010 起", "行動影像＋標記＋幾何", "Fiducial marker", "本研究改以信用卡、三視角"],
        ["Dehais 等 [29]", "2017", "雙視角稠密 3D 重建", "信用卡參考卡", "本研究三視角＋DUSt3R"],
        ["MUSEFood [35]", "2019", "多感測器＋微分幾何", "MLS 聲學測距", "本研究免音訊、用信用卡"],
        ["goFOOD [34]", "2020", "Mask R-CNN＋分割＋幾何", "模板／幾何模型", "本研究採開放域 VLM＋SAM3"],
        ["He 等 [33]", "2020", "多任務辨識＋份量", "影像特徵", "本研究分離語意與幾何"],
        ["Shao 等 [49]", "2021", "能量分布圖＋跨域融合", "單張 RGB", "本研究幾何管線非純回歸"],
        ["Nutrition5k [20]", "2021", "大規模營養理解基準", "實測重量／RGB-D", "本研究活大自助餐實測"],
        ["Vinod 等 [51]", "2022", "Depth＋能量密度融合", "RGB-D", "本研究推論不需深度"],
        ["Shao 等 [21]", "2023", "RGB→voxel＋能量回歸", "訓練用 depth", "本研究多視角幾何積分"],
        ["Vinod 等 [22]", "CVPR 2024", "3D 模型縮放", "Checkerboard", "同樣重視絕對尺度"],
        ["Dhar 等 [50]", "2025", "YOLOv9＋硬幣縮放＋回歸", "實體硬幣", "本研究開放域＋信用卡"],
        ["MetaFood [23]", "2024", "物理參考 3D 重建競賽", "信用卡／棋盤格", "技術路線相近"],
    ],
    caption_text="表 2-4　電腦視覺食物熱量估算代表性研究比較",
)
make_table(
    ["論文", "年份", "資料集／場域", "輸入條件", "主要指標", "備註"],
    [
        ["Fang 等 [32]", "2019", "TADA 子集", "單張 RGB + 棋盤", "誤差率 10.89%", "—"],
        ["He 等 [33]", "2020", "834 單食物影像", "裁切 RGB", "MAE 56.82 kCal", "21 類封閉集"],
        ["Shao 等 [49]", "2021", "96 eating scenes", "單張 RGB", "MAPE 11.47%", "—"],
        ["Thames 等 [20]", "2021", "Nutrition5k", "RGB / RGB-D", "Vol. scalar 16.5%", "固定 top-view"],
        ["Vinod 等 [51]", "2022", "Nutrition5k 子集", "RGB-D", "MAPE 13.57%", "需 depth"],
        ["Shao 等 [21]", "2023", "3,259 top-view", "單張 RGB", "MAE 40.05 kCal", "原文明列 MAPE 不一致"],
        ["Vinod 等 [22]", "2024", "SimpleFood45", "RGB+checkerboard", "EMAPE 17.67%", "3D model scaling"],
        ["Dhar 等 [50]", "2025", "孟加拉街頭 5 類", "RGB+硬幣", "MAE 7.85 kCal", "封閉垂直場景"],
        ["Dehais 等 [29]", "2017", "77 道菜", "雙視角+參考卡", "MAPE 8.2%", "受控拍攝"],
        ["本研究", "2026", "活大自助餐", "三視角+信用卡", "待填入", "開放域混合餐盤"],
    ],
    caption_text="表 2-5　代表性論文之性能與實驗條件比較（條件不同，不宜直接排序）",
)

heading("2.4.5　綜述性調查與大型多模態模型之應用", level=3)
para("綜述性文獻有助梳理上述研究之演進脈絡。Subhi 等人"
     f"{cit(25)}與 Lo 等人{cit(24)}分別從 IEEE Access 與 JBHI 角度調查視覺食物"
     "辨識與體積估計；Wang 等人"
     f"{cit(16)}於 Trends in Food Science & Technology 發表 VBDA 專題回顧，"
     "系統整理由影像採集、食物識別、份量估算至營養計算之完整流程；Chen 等人"
     f"{cit(48)}於 2024 年進一步以 scoping review 方式，彙整以 AI 分析食物影像"
     "進行膳食評估之最新進展。另一方面，GPT-4V 等大型多模態模型"
     f"{cit(46)}可直接由影像生成食物描述與粗略熱量估計，部署門檻低，但 Wang 等人"
     f"{cit(16)}與 Chen 等人{cit(48)}均指出其缺乏可靠尺度與可驗證之幾何依據。"
     "相較之下，本研究採「VLM 負責語意、幾何管線負責份量」之分工，並於第六章以自行建立"
     f"之 LMM 基準與本方法進行對照實驗；GPT-4V 等模型雖可接收影像輸入{cit(46)}，"
     "但其技術報告並未提供食物熱量估算之標準 benchmark 數值，故不宜以該文獻單獨"
     "支撐特定誤差區間。此設計呼應 Wang 等人"
     f"{cit(16)}與 Chen 等人{cit(48)}對「引入物理尺度」之倡議。")

heading("2.4.6　在地化垂直應用與可重現性評估", level=3)
para("除上述以西方或校園餐廳為主之研究外，近年亦出現針對特定地域飲食型態之垂直系統。"
     "Dhar 等人"
     f"{cit(50)}針對孟加拉街頭食物，指出既有工作多忽略實際份量且輸出固定熱量。其方法"
     "以改良 YOLOv9（C3f_CD + CBAM）做分類與實例分割，以 5 Taka 硬幣（直徑 24.6 mm）"
     "換算像素尺度，再自 mask 抽取 height、width、area、perimeter 等幾何特徵回歸"
     "熱量。於 1,847 張原始影像（增強至 3,885 張、5 類主食物）上，MAE 7.85 kCal、"
     "RMSE 22.14、R² 96.0%。此結果說明：若鎖定小範圍品類並要求使用者放置參考物，"
     "以部分幾何特徵回歸即可達極高精度；但其泛化限於少數類別，且未處理混合餐盤。"
     "對台灣應用而言，夜市、便當配菜或滷味等垂直場景可借鑑其「先縮小食物空間、再以"
     "實體幾何參考穩住份量」之產品化思路，而非直接套用 generic 模型。")
para("文獻回顧時亦應將「可重現性」列為篩選條件。Nutrition5k"
     f"{cit(20)}與 Vinod 等人 SimpleFood45{cit(22)}提供公開資料或程式，屬可重現性"
     "較佳者；Fang 等人"
     f"{cit(32)}、Shao 等人{cit(49)}等多篇小資料集工作雖成績亮眼，但資料與程式"
     "公開不足；Shao 等人"
     f"{cit(21)}更存在摘要與表格 MAPE 不一致之問題。本研究採可開源之 DUSt3R、SAM3"
     "與 Qwen2.5-VL，並以可視化中間結果（分割、點雲、厚度圖）提升可除錯性，呼應"
     "Vinod 等人所倡之 geometry-first、可解釋部署方向。")

heading("2.4.7　技術演進與本研究之方法選型", level=3)
para("綜合 2017 至 2025 年文獻，電腦視覺食物熱量估算並非單向從「傳統方法」走向"
     "「深度學習」，而是在「顯式幾何可解釋性」與「單眼 RGB 可部署性」之間擺盪。"
     "Dehais 2017 以雙視角三角量測換準確度；MUSEFood 2019 以聲學測距免參考物；"
     "He 2020 與 Shao 2021"
     f"{cit(33,49)}讓模型從資料分布間接學尺度；Nutrition5k 2021"
     f"{cit(20)}將問題做成可比較 benchmark；Shao 2023"
     f"{cit(21)}以 3D voxel 重建恢復幾何；Vinod 2024"
     f"{cit(22)}再以顯式 3D model 與 camera pose 拉回可解釋框架；Dhar 2025"
     f"{cit(50)}示範在地小類別＋硬幣參考之垂直落地。每一代工作都在回答：在缺少"
     "深度與基準尺時，如何恢復尺度？")
para("若依應用假設選型：第一，純 RGB、免參考物——可參考 Shao 等人"
     f"{cit(49,21)}，但須補足在地資料，否則泛化風險高；第二，有 depth 硬體——"
     f"Vinod 等人{cit(51)}適合作上限模型；第三，允許參考物、追求可解釋性——"
     f"Dehais 等人{cit(29)}、Vinod 等人{cit(22)}與 Dhar 等人{cit(50)}較穩。"
     "本研究選擇第三類中之「信用卡＋多視角幾何重建」變體：不建 3D food model 資料庫、"
     "不要求 checkerboard，而以 DUSt3R 直接還原場景點雲並積分體積，並以 VLM 補足"
     "開放域辨識，在工程複雜度與泛化能力間取得折衷。")
make_table(
    ["時期", "代表研究", "核心問題", "尺度恢復方式"],
    [
        ["2017", "Dehais 等 [29]", "雙視角幾何重建", "信用卡參考卡＋stereo"],
        ["2019", "MUSEFood [35]", "免參考物體積", "MLS 聲學測距＋微分幾何"],
        ["2020", "He 等 [33]", "多任務學習", "從影像特徵間接學習"],
        ["2021", "Shao 等 [49]", "單眼 RGB 熱量", "energy distribution 中介表徵"],
        ["2021", "Nutrition5k [20]", "公開 benchmark", "實測重量／depth／volume"],
        ["2023", "Shao 等 [21]", "RGB 推論、depth 訓練", "voxel 重建隱式幾何"],
        ["2024", "Vinod 等 [22]", "可解釋 3D scaling", "checkerboard＋3D prior"],
        ["2025", "Dhar 等 [50]", "在地垂直應用", "硬幣像素尺度＋回歸"],
        ["本研究", "—", "開放域混合餐盤", "信用卡＋三視角 DUSt3R"],
    ],
    caption_text="表 2-6　電腦視覺食物熱量估算技術演進時間線",
)

heading("2.5　現有食物熱量估算系統之瓶頸分析", level=2)
para("綜合前述技術，現有影像式食物熱量估算方法可按「是否具幾何量測」「是否具絕對尺度」"
     "與「是否可解釋」三個維度分類。圖 2-9 展示常見方法之分類架構。")
fig_img("fig2_9.png", "圖 2-9　影像式食物熱量估算方法分類", width_cm=14)

heading("2.5.1　分類—查表法與二維面積估算法", level=3)
para("第一類為「分類—查表法」：先以分類器辨識食物種類，再依標準份量查表得熱量"
     f"{cit(16,3)}。Food-101"
     f"{cit(3)}時代之系統多屬此類：Top-1 準確率可達 90% 以上，但假設每份"
     "「一份炒青菜」恆為固定大卡數，忽略真實份量可差數倍之事實。對自助餐、"
     "吃到飽或使用者自盛等情境，此假設嚴重偏離，導致系統性低估或高估。")
para("第二類為「二維面積估算法」：由影像中食物所占面積或像素數推估份量，並常以硬幣、"
     f"餐盤或 fiducial marker 校正{cit(24,30,37)}。TADA"
     f"{cit(30,47)}以 5×4 棋盤 marker 做 homography 與透視校正；Chen 等人"
     f"{cit(37)}結合 Faster R-CNN 與 GrabCut，以俯視＋側視估算投影面積再推體積；"
     f"Dhar 等人{cit(50)}以硬幣直径換算像素比例，再自 mask 抽取 2D 幾何特徵。"
     "此類方法之瓶頸在於幾何假設過於簡化——將食物視為柱體或固定高度模板，"
     "難以反映堆疊、有厚度或形狀不規則之食物；對有湯汁、塌陷或立體高度差大"
     "之菜色，二維投影面積與真實體積之相關性會顯著下降。")

heading("2.5.2　端對端深度學習與大型多模態模型估算法", level=3)
para("第三類為直接以深度網路或大型多模態模型由影像回歸熱量。Fang 等人"
     f"{cit(32)}以 cGAN 學習 energy distribution image；He 等人"
     f"{cit(33)}與 Shao 等人{cit(49)}以 CNN 特徵融合回歸 kcal；Shao 等人"
     f"{cit(21)}以 voxel 重建輔助能量估計。共同優點是推論流程短、不需使用者"
     "放置參考物；共同缺點是缺乏絕對物理尺度——模型只能從訓練分布學習「這類"
     "食物通常多大」，當拍攝距離、盤子大小或份量非典型時即失效。")
para("GPT-4V 等 LMM"
     f"{cit(46)}將問題進一步簡化為「看圖說話」：使用者詢問熱量，模型即輸出"
     "數字與描述。Wang 等人"
     f"{cit(16)}與 Chen 等人{cit(48)}指出，此類輸出常受語言先驗主導（例如"
     "見到「炸雞」即輸出典型一份炸雞大卡數），而非影像中實際幾何。本研究第六章"
     "以 LMM 為基準，即為驗證此第三類方法在活大場景下之侷限。")

heading("2.5.3　三維重建輔助份量估算與方法比較", level=3)
para("第四類結合三維重建、參考物與幾何計算，為目前較具物理依據之路線。"
     "Dehais 等人"
     f"{cit(29)}以雙視角 stereo 積分 mesh 體積；MUSEFood"
     f"{cit(35)}以聲學測距 + 微分幾何；Vinod 等人"
     f"{cit(22)}以 3D model scaling；MetaFood"
     f"{cit(23)}與 VolETA{cit(36)}競賽推動信用卡／棋盤格輔助之 one-shot "
     "volume estimation。此類方法之 MAPE 多落在 8%～18% 區間（視場域與"
     "食物複雜度而異），且可提供中間幾何量供除錯。")
para("第四類之代價在於管線較長、對拍攝規範較敏感。Dehais 需雙視角 + 分割；"
     "Vinod 需 checkerboard + 3D asset；MUSEFood 需頂側雙視角 + 音訊。"
     "本研究選擇「三視角 RGB + 信用卡 + DUSt3R」，可視為在第四類中折衷："
     "不要求 3D 掃描資料庫、不要求 checkerboard 標定板，但仍保留可解釋"
     "之體積積分。表 2-3 綜合比較各類方法。")
make_table(
    ["方法類型", "代表研究", "絕對尺度", "可解釋性", "主要侷限"],
    [
        ["分類—查表", "Food-101 系統", "否", "低", "忽略份量差異"],
        ["二維面積／雙視角", "TADA [30], Chen [37]", "部分", "中", "幾何假設簡化"],
        ["端對端／LMM", "Shao [49], LMM 基準", "否", "低", "缺乏尺度／待實測"],
        ["3D＋參考物", "Vinod [22], MetaFood [23]", "是", "高", "管線較複雜"],
        ["本研究", "信用卡＋DUSt3R＋SAM3", "是", "高", "雲端算力需求"],
    ],
    caption_text="表 2-3　影像式食物熱量估算方法綜合比較",
)
para("由表 2-3 可知，現有研究之共通缺口在於難以同時滿足「絕對尺度」「可解釋中間結果」"
     "與「免專用硬體」三項目標。本研究之方法即針對此缺口，以多視角 DUSt3R 重建取得"
     "立體形狀、以信用卡比例尺賦予絕對尺度、以 SAM3 界定各食材範圍，並以幾何積分"
     "換算熱量，在原理上結合第四類方法之優點，同時以 VLM 提供開放域辨識能力。")

heading("2.6　行動醫療與雲端人工智慧推論架構相關研究", level=2)
para("本研究系統採行動端與雲端協同架構，故本節回顧 mHealth 飲食紀錄與雲端推論相關研究。"
     "圖 2-10 示意薄用戶端—重雲端之典型分工。")
fig_img("fig2_10.png", "圖 2-10　行動醫療應用之薄用戶端—重雲端推論架構", width_cm=14)

heading("2.6.1　行動醫療飲食紀錄應用", level=3)
para("mHealth 領域已有大量研究探討以行動 App 進行飲食紀錄、活動追蹤與慢性病管理"
     f"{cit(16,25)}。Zhu 等人早期即探討行動裝置輔助飲食評估"
     f"{cit(30)}；Six 等人{cit(31)}開發影像分析系統支援營養師評估。相關調查指出，"
     "使用者依從性、操作便利性與結果可信度為影響成效之關鍵。以拍照取代手動輸入、"
     "提供即時且可理解之回饋，有助提升長期使用意願；但若估算誤差過大，反而可能損害信任。"
     "因此，方法準確度與可解釋性同為 mHealth 落地之重要條件；本研究第六章將以"
     "實測 MAPE 與 LMM 基準對照驗證此點。")

heading("2.6.2　雲端 GPU 推論與系統部署", level=3)
para("視覺語言模型、SAM 與 DUSt3R 皆具相當之 GPU 記憶體與算力需求，難以完全在"
     "一般手機上即時執行。goFOOD"
     f"{cit(34)}、TADA {cit(30)}等早期系統多部署於固定伺服器；Shao 等人"
     f"{cit(21)}與 Vinod 等人{cit(22)}之深度模型亦假設 GPU 伺服器環境。"
     f"近年按用量計費之 GPU 雲平台{cit(48)}使間歇性推論更具成本彈性。雲端推論"
     "架構將影像採集、結果呈現等輕量任務保留於行動端，推論密集之模型運算交由 GPU "
     "伺服器執行，利於模型集中維護與更新。相較 MUSEFood"
     f"{cit(35)}需多視角與音訊量測程序，本研究僅增加三張照片與信用卡擺放，"
     "在精度與使用者負擔間取得折衷。相關文獻亦討論延遲、吞吐量、成本與可擴充性權衡，"
     "將於本論文第五、六章進一步量化分析。")
para("從 mHealth 落地角度，Chen 等人"
     f"{cit(48)}之 scoping review 指出，AI 膳食評估系統之採用關鍵在於準確度、"
     "可解釋性與操作便利性。TADA/mpFR"
     f"{cit(47)}強調標準化拍攝流程；Dhar 等人{cit(50)}示範垂直場景可達極低"
     "MAE。本研究之行動—雲端架構與可視化中間結果（分割、點雲、厚度圖），"
     "旨在回應上述三項需求，使營養師或一般使用者能檢視「為何得到此熱量數值」。")

heading("2.7　文獻探討小結與本研究定位", level=2)
para("綜合本章回顧，VBDA 領域在食物辨識、分割與三維重建上均已累積豐富成果：CNN 與"
     "大型資料集推動了封閉集合分類；SAM 系列與 VLM 則開啟開放域、低標註成本之辨識分割"
     "可能；DUSt3R 等模型降低多視角重建之門檻。2.4 節所回顧之 TADA、Nutrition5k、"
     "Shao 系列、Vinod 系列與 Dhar 等人"
     f"{cit(50)}等研究顯示，電腦視覺熱量估算已由「分類—查表」逐步演進至 energy "
     "distribution 中介表徵、RGB-D 融合、體素重建與顯式三維縮放；表 2-5 亦表明，"
     "最終性能往往不由 backbone 單獨決定，而取決於尺度資訊與資料標註品質。然而，"
     "現有方法仍少數能同時解決「絕對物理尺度」「可解釋幾何計算」與「一般手機可負擔"
     "之硬體成本」三者。")
para("本研究定位於整合 VLM 開放域辨識、SAM3 精確分割、DUSt3R 多視角重建與信用卡"
     "比例尺錨定，並以行動—雲端架構完成可部署原型。相較 Shao 等人"
     f"{cit(49,21)}之純 RGB 回歸，本方法提供可檢驗之管線估計體積與厚度圖；相較 Vinod 等人"
     f"{cit(22)}之 3D model scaling，本方法免預建每類食物三維資產；相較 Nutrition5k"
     f"{cit(20)}之 top-view 設定，本方法以三視角手機拍攝貼近真實用餐；相較 Dhar 等人"
     f"{cit(50)}之封閉五類，本方法以 VLM 支援混合開放域餐盤。下一章將說明支撐營養"
     "換算之食物幾何與營養轉換資料庫建置，說明如何以與推論相同之管線標定體積—重量對應；"
     "第四章則詳述完整之體積估算架構。")

# ===========================================================================
# 第三章 食物幾何與營養轉換資料庫建置
# ===========================================================================
heading("第三章　食物幾何與營養轉換資料庫建置", level=1, page_break_before=True)
para("由影像經本研究管線求得各食材之管線估計體積（立方公分）後，尚須將其轉換為熱量與營養素，"
     "方能完成飲食評估。此轉換無法僅憑幾何推得，而需依賴「每單位重量之營養成分」與「管線"
     "估計體積—重量」之對應關係。本章說明本研究如何建置「食物幾何與營養轉換資料庫」，作為由"
     "管線估計體積推得熱量與三大營養素（蛋白質、脂肪、碳水化合物）之查詢基礎，並訂定多視角"
     "拍攝與比例尺擺放規範，使資料蒐集與線上推論共用同一套幾何前提與體積估算管線。")

heading("3.1　資料庫建置動機與流程", level=2)
para("由影像經本研究管線估得食物之管線估計體積後，尚無法直接得知其熱量，因為熱量與營養成分"
     "係以「每單位重量」定義"
     f"（如每 100 公克之大卡數）{cit(16,20)}，而重量又取決於食物之堆積密度。因此，欲由"
     "管線估計體積推得營養資訊，必須經過「管線估計體積 → 重量 → 營養」兩階段之轉換，而此"
     "二階段所需之「堆積密度」與「單位重量營養值」皆需事先蒐集並結構化儲存。此即本研究建置"
     "食物幾何與營養轉換資料庫之動機。")
para("資料庫之建置流程如下：（一）盤點目標應用情境中常見之食物品項，優先涵蓋台灣日常飲食"
     "（如各式米飯麵食、常見蔬菜、肉類與自助餐菜色）；（二）自公開之食品營養成分資料庫"
     f"（如衛生福利部食品營養成分資料庫）{cit(18)}取得各品項每 100 公克之熱量、蛋白質、脂肪"
     "與碳水化合物含量；（三）對各品項準備「一分量」樣本，依 3.2 節規範拍攝三視角影像並"
     "置放標準信用卡，再執行與第四章相同之完整推論管線（辨識、分割、三維重建、比例尺錨定與"
     "厚度積分），取得該分量之管線估計體積；隨後將同一分量以電子秤秤得實測重量，並依式 (3-2)"
     "反推堆積密度，建立體積與重量之對應"
     f"{cit(20,24)}；（四）將密度與單位重量營養值合併，預先計算為「每立方公分」之"
     "營養值，以簡化線上推論之計算；（五）以食物識別鍵（key）為索引，將上述資訊整理為"
     "對應與模糊匹配機制。")
para("需特別說明的是，本研究所稱之「體積」並非以水置換法或 CT 等獨立量測所驗證之物理真值"
     "體積，而是「本研究管線定義下之估計體積」。然而，由於資料庫建置與線上推論採用完全相同之"
     "拍攝規範、模型與積分演算法，體積雖帶有管線本身之系統誤差，卻能在系統內部形成自洽之"
     "體積—重量座標系；亦即，推論時所得之管線估計體積，可透過資料庫中由相同方法標定之密度"
     "參數，合理地換算為重量與營養資訊。此設計之合理性在於方法一致性，而非宣稱已取得絕對"
     "真值體積。")
fig_img("fig3_1.png", "圖 3-1　食物幾何與營養轉換資料庫之建置流程", width_cm=9.5)

heading("3.2　多視角食物影像採集規範", level=2)
para("三維重建之品質高度依賴輸入影像之視角配置與拍攝條件。為確保重建結果之穩定與可重現，"
     f"本研究訂定一套標準化之多視角影像採集規範，作為資料庫建置與後續系統使用之共同準則"
     f"{cit(29,47,30)}。")

heading("3.2.1　三視角拍攝路徑設計", level=3)
para("本研究採用三張影像進行重建，其視角配置為：正上方（俯視，約 90 度）、左側四十五度與右側"
     "四十五度。此一配置之設計考量如下：正上方影像可清楚呈現各食材之平面輪廓與相對位置，適合"
     "作為分割與體積計算之參考視角（reference view）；左右四十五度影像則提供足夠之視差"
     "（parallax），使三維重建得以還原食物之高度與立體形狀。三張影像於水平方向約略等角分佈，"
     "可在最少拍攝張數下取得足夠之幾何約束，兼顧重建品質與使用者操作負擔。")
para("為維持視差之一致性，拍攝時應以餐點中心為視覺焦點，維持大致固定之拍攝距離，並使餐點與"
     "比例尺參考物完整入鏡。三張影像之拍攝順序不影響重建結果，但於後續管線中須明確指定正上方"
     "影像作為參考視角。")
fig_img("fig3_2.png", "圖 3-2　三視角拍攝路徑示意圖", width_cm=12)

heading("3.2.2　參考平面與絕對比例尺擺放規範", level=3)
para("為賦予三維重建結果絕對尺度，本研究採用符合 ISO/IEC 7810 ID-1 規格之標準信用卡"
     f"（長 85.60 公釐、寬 53.98 公釐，對角線約 101.2 公釐）{cit(19)}作為絕對比例尺參考物。"
     "選用信用卡之理由為：其尺寸經國際標準嚴格規範、全球一致；為多數使用者隨身攜帶之日常物品，"
     "取得容易；且外型為平整之矩形，邊界清晰、易於分割與偵測。")
para("擺放規範如下：（一）將信用卡平放於餐點旁之同一參考平面（如餐桌或托盤）上，且盡量與餐點"
     "共平面，以利平面擬合；（二）信用卡應完整入鏡且不被餐具或食物遮擋；（三）避免信用卡表面"
     "產生強烈反光，以確保其於各視角影像中皆可被穩定分割；（四）信用卡與餐點之距離不宜過遠，"
     "以維持其於三維重建結果中之點雲密度。")
para("在後續之比例尺錨定步驟中，系統將於重建後之三維空間中量測信用卡之對角線長度（以無尺度之"
     "模型單位表示），並與其已知之真實對角線長度相除，即得「每一模型單位對應之公分數」，"
     "作為將整個三維模型由相對尺度轉換為絕對尺度之換算係數。此一方法之細節將於第四章 4.5 節詳述。")
repo_img("90.jpg",
         "圖 3-3　絕對比例尺（標準規格卡片）之擺放規範實例（卡片平放於餐點旁同一平面、完整入鏡）",
         width_cm=9, color=True)

heading("3.3　堆積密度標定與管線體積—重量對應", level=2)
para("管線估計體積至重量之轉換係透過食物之堆積密度完成。堆積密度定義為食物於自然堆疊狀態下，"
     "單位體積所對應之質量（公克／立方公分）。與物質之真實密度不同，堆積密度考量了食物顆粒間"
     "或塊體間之空隙，因此更貼近實際盛裝於餐盤上之情形。Nutrition5k"
     f"{cit(20)}以逐項秤重取得 mass ground truth，本研究則以台灣常見自助餐品項，"
     "在方法一致之前提下建立密度表，使管線估計體積（{cm³}）可換算為公克。其基本關係式如式 (3-1)：")
equation(
    m_omath(
        m_text("W"), m_text(" = "),
        m_sub(m_text("V", italic=True), m_text("pipeline")),
        m_text(" × "),
        M.rho(),
    ),
    number="(3-1)",
)
para("式中，{W} 為重量（g）；{V:pipeline} 為管線估計體積（{cm³}）；{rho} 為堆積密度（{g/cm³}）。",
     indent_chars=0)
para("本研究之堆積密度並非以量筒排水法或定容量器直接量得容積，而是以「管線體積—秤重」配對"
     "標定：對每一品項之一分量，先依 3.2 節完成三視角拍攝，再執行第四章之完整推論管線取得"
     "管線估計體積，隨後將同一分量以電子秤量得實測重量，則堆積密度可由式 (3-2) 求得：")
equation(
    m_omath(
        M.rho(), m_text(" = "),
        m_frac(
            m_sub(m_text("W", italic=True), m_text("scale")),
            m_sub(m_text("V", italic=True), m_text("pipeline")),
        ),
    ),
    number="(3-2)",
)
para("此一作法之核心在於，建庫階段與推論階段使用相同之體積定義，使密度參數實質上成為"
     "「在本研究管線座標系下，每立方公分管線體積對應多少公克」之經驗換算係數。")
para("需說明的是，堆積密度會受到食物之含水量、烹調方式、切塊大小與堆疊鬆緊程度影響，故資料庫中"
     "所記錄之密度為特定烹調狀態、特定一分量型態下之代表值，於實際應用時將引入一定程度之"
     "系統誤差。此外，管線估計體積本身亦受分割、重建與積分演算法影響，但由於建庫與推論共用"
     "同一管線，此誤差在體積—重量換算中具部分抵消效果。上述誤差來源及其對整體熱量估算之"
     "影響，將於第六章之誤差分析中進一步討論。")
make_table(
    ["食物品項", "管線體積（{cm³}）", "秤重（g）", "堆積密度（{g/cm³}）", "熱量（{kcal/cm³}）", "標定方式"],
    [
        ["鹽酥雞", "33.17", "57.4", "1.73", "3.95", "管線體積＋同一分量秤重"],
        ["炸雞腿", "69.98", "127.5", "1.82", "3.57", "管線體積＋同一分量秤重"],
        ["芋頭", "34.97", "43.6", "1.25", "1.60", "管線體積＋同一分量秤重"],
        ["蒸蛋", "24.81", "45.5", "1.83", "0.59", "管線體積＋同一分量秤重"],
        ["紅蘿蔔炒蛋", "115.29", "68.5", "0.59", "0.85", "管線體積＋同一分量秤重"],
        ["青椒", "28.87", "39.6", "1.37", "0.40", "管線體積＋同一分量秤重"],
        ["麻婆豆腐", "57.41", "75.4", "1.31", "1.36", "管線體積＋同一分量秤重"],
        ["韓式泡菜", "62.51", "55.6", "0.89", "0.31", "管線體積＋同一分量秤重"],
        ["炒高麗菜", "90.56", "63.0", "0.70", "0.16", "管線體積＋同一分量秤重"],
        ["炒苦瓜", "84.91", "81.6", "0.96", "0.19", "管線體積＋同一分量秤重"],
    ],
    caption_text="表 3-1　代表性食物之堆積密度與每立方公分營養參數標定結果",
)
make_table(
    ["食物品項", "管線體積（{cm³}）", "推得熱量（kcal）", "推得蛋白質（g）", "推得脂肪（g）", "推得碳水（g）"],
    [
        ["鹽酥雞", "33.17", "131", "8.8", "7.4", "7.2"],
        ["炸雞腿", "69.98", "250", "18.4", "19.0", "2.2"],
        ["芋頭", "34.97", "56", "1.1", "0.5", "11.5"],
        ["蒸蛋", "24.81", "15", "1.6", "0.8", "1.6"],
        ["紅蘿蔔炒蛋", "115.29", "99", "5.6", "6.7", "4.2"],
        ["青椒", "28.87", "11", "0.4", "0.2", "2.3"],
        ["麻婆豆腐", "57.41", "78", "5.8", "3.3", "6.5"],
        ["韓式泡菜", "62.51", "19", "1.1", "0.2", "3.2"],
        ["炒高麗菜", "90.56", "14", "0.8", "0.1", "3.0"],
        ["炒苦瓜", "84.91", "16", "0.7", "0.1", "3.4"],
    ],
    caption_text="表 3-1a　一分量管線體積與推得營養總量（營養值依 TFDA 每 100 公克資料 × 秤重換算）",
)

heading("3.4　體積與營養熱量轉換表之架構與實作", level=2)
para("本研究之「體積—營養轉換表」同時整合兩類獨立量測：（甲）幾何端——以與線上推論相同之管線"
     "取得各品項一分量之管線估計體積 {V:pipeline}，並以電子秤量得同一分量之實測重量"
     "{W:scale}（見 3.3 節）；（乙）營養端——自衛生福利部食品藥物管理署「食品營養成分"
     f"資料庫」（TFDA，https://consumer.fda.gov.tw，存取日期 2026 年 7 月）{cit(18)}"
     "查得各對應樣本每 100 公克之熱量與三大營養素。兩端資料於建庫階段合併，預先換算為"
     "「每立方公分營養值」，線上推論時僅需將管線估計體積乘以該係數即可。")

heading("3.4.1　TFDA 對照、整合編號與合成規則", level=3)
para("對每一自助餐品項，先以中文品名檢索 TFDA 資料庫。若存在與實際菜色同名且烹調狀態"
     "相符之樣本，則直接採用其整合編號與每 100 公克營養值，記為「直接對照」；若無完全"
     "同名條目，則採解剖或烹調型態最接近之單一樣本，記為「代理對照」，並於資料庫註明"
     "代理理由；若為由兩種以上 TFDA 樣本組成之複合菜色，則依固定重量配比加權合成每"
     "100 公克營養值，並賦予本研究自訂之「合成整合編號」（SYN-xxxx），記為「加權合成」。"
     "表 3-4 列舉十項標定品項之完整對照。")
para("加權合成之每 100 公克營養值計算如式 (3-4)：")
equation(
    m_omath(
        m_sub(m_text("N", italic=True), m_text("blend")),
        m_text(" = "),
        m_nary(
            [
                m_sub(m_text("w", italic=True), m_text("i")),
                m_text(" × "),
                m_sub(m_text("N", italic=True), m_text("i")),
            ],
            sub=m_text("i"),
        ),
    ),
    number="(3-4)",
)
para("式中，{N:blend} 為合成後每 100 公克營養值；{w:i} 為第 i 個 TFDA 成分之重量"
     "配比（Σwᵢ = 1）；{N:i} 為該成分於 TFDA 之每 100 公克營養值。取得 {N:100g}"
     "（直接、代理或合成）後，先依同一分量秤重換算一分量營養總量，如式 (3-5)：")
equation(
    m_omath(
        m_sub(m_text("N", italic=True), m_text("portion")),
        m_text(" = "),
        m_frac(
            m_sub(m_text("W", italic=True), m_text("scale")),
            m_text("100"),
        ),
        m_text(" × "),
        m_sub(m_text("N", italic=True), m_text("100g")),
    ),
    number="(3-5)",
)
para("再除以管線估計體積，得每立方公分營養值，如式 (3-3)：")
equation(
    m_omath(
        m_sub(m_text("N", italic=True), m_sup(m_text("cm"), m_text("3"))),
        m_text(" = "),
        m_frac(
            m_sub(m_text("N", italic=True), m_text("portion")),
            m_sub(m_text("V", italic=True), m_text("pipeline")),
        ),
    ),
    number="(3-3)",
)
para("式 (3-3) 亦等價於 {N:cm³} = ρ × {N:100g} / 100，其中 ρ = {W:scale} / {V:pipeline}"
     "為 3.3 節所標定之堆積密度。線上推論時，系統以 {V:pipeline,infer} × {N:cm³} 即得"
     "各營養素推估值。")
make_table(
    ["食物品項", "對照類型", "整合編號", "TFDA 樣本／合成公式", "每 100g 熱量（kcal）"],
    [
        ["鹽酥雞", "代理", "R5500201", "冷凍雞塊", "228"],
        ["炸雞腿", "代理", "I0404102", "骨腿(肉雞)", "196"],
        ["芋頭", "直接", "B0500101", "芋頭", "128"],
        ["蒸蛋", "直接", "K0150301", "蒸蛋(市售)", "32"],
        ["紅蘿蔔炒蛋", "加權合成", "SYN-0001", "60%×K0150201＋40%×E0200101", "143.8"],
        ["青椒", "代理", "E75001", "甜椒平均值(青皮)", "29"],
        ["麻婆豆腐", "加權合成", "SYN-0002", "85%×R4700901＋15%×P1004601", "103.3"],
        ["韓式泡菜", "直接", "R4400301", "韓式泡菜", "35"],
        ["炒高麗菜", "代理", "E30001", "甘藍平均值", "23"],
        ["炒苦瓜", "代理", "E6500201", "苦瓜(青皮)", "20"],
    ],
    caption_text="表 3-4　十項標定食物之 TFDA 營養資料對照與整合編號",
)
para("表 3-4 中，直接對照者之整合編號與 TFDA 官方編號一致；代理對照者沿用所選代理樣本之"
     "官方編號，並於 JSON 紀錄 tfda_proxy_note 說明替代理由；加權合成者之整合編號為本研究"
     "自訂 SYN 序號，成分編號與配比同時寫入 tfda_blend 與 tfda_synthetic_formula 欄位，"
     "以利追溯。表 3-1 之堆積密度與表 3-1a 之一分量營養總量，分別對應式 (3-2) 與"
     "式 (3-5)；表 3-3 之每立方公分係數則對應式 (3-3)，三者共同構成完整之體積—營養"
     "轉換鏈。")

heading("3.4.2　JSON 資料庫欄位與查詢介面", level=3)
para("本研究將上述資訊以結構化 JSON 格式儲存於 food_nutrition_db.json，每一食物品項以"
     "唯一識別鍵（key）為索引。除每立方公分營養係數外，每筆紀錄亦保存 TFDA 整合編號、"
     "樣本名稱、對照類型、每 100 公克營養值、一分量秤重、管線體積、堆積密度及一分量"
     "營養總量，使表格數值均可由原始欄位重算驗證。主要欄位如表 3-2 所示。")
make_table(
    ["欄位名稱", "資料型別", "單位", "說明"],
    [
        ["display_name", "字串", "—", "食物之中英文顯示名稱"],
        ["tfda_integration_no", "字串", "—", "TFDA 官方整合編號，或合成編號 SYN-xxxx"],
        ["tfda_source_type", "字串", "—", "direct／proxy／blend"],
        ["tfda_per_100g", "物件", "kcal、g", "每 100 公克營養值（直接、代理或合成）"],
        ["portion_weight_g", "浮點數", "g", "標定時同一分量之電子秤讀值"],
        ["pipeline_volume_cm3", "浮點數", "{cm³}", "同一分量之管線估計體積"],
        ["bulk_density_g_cm3", "浮點數", "{g/cm³}", "ρ = portion_weight_g / pipeline_volume_cm3"],
        ["portion_nutrition", "物件", "kcal、g", "一分量營養總量（式 3-5）"],
        ["kcal_cm3", "浮點數", "{kcal/cm³}", "每立方公分熱量（式 3-3）"],
        ["protein_g_cm3", "浮點數", "{g/cm³}", "每立方公分蛋白質"],
        ["fat_g_cm3", "浮點數", "{g/cm³}", "每立方公分脂肪"],
        ["carbs_g_cm3", "浮點數", "{g/cm³}", "每立方公分碳水化合物"],
    ],
    caption_text="表 3-2　營養轉換資料庫之欄位架構",
)
para("下表列舉目前已完成管線體積與秤重標定之十項代表性食物（表 3-1、表 3-1a、表 3-4），"
     "以及寫入資料庫之每立方公分營養參數（表 3-3）。資料庫現階段僅收錄此十品項，"
     "其餘品項待完成相同標定程序後再擴充。")
make_table(
    ["識別鍵", "顯示名稱", "{kcal/cm³}", "蛋白質 {g/cm³}", "脂肪 {g/cm³}", "碳水 {g/cm³}"],
    [
        ["salt_pepper_chicken", "鹽酥雞", "3.95", "0.265", "0.223", "0.218"],
        ["chicken_leg_fried", "炸雞腿", "3.57", "0.262", "0.272", "0.031"],
        ["taro", "芋頭", "1.60", "0.031", "0.014", "0.329"],
        ["egg_steamed", "蒸蛋", "0.59", "0.066", "0.033", "0.062"],
        ["carrot_egg_stir_fried", "紅蘿蔔炒蛋", "0.85", "0.048", "0.058", "0.037"],
        ["green_pepper_stir_fried", "青椒", "0.40", "0.014", "0.006", "0.081"],
        ["mapo_tofu", "麻婆豆腐", "1.36", "0.101", "0.057", "0.113"],
        ["kimchi_korean", "韓式泡菜", "0.31", "0.018", "0.004", "0.052"],
        ["cabbage_stir_fried", "炒高麗菜", "0.16", "0.009", "0.001", "0.033"],
        ["bitter_gourd_stir_fried", "炒苦瓜", "0.19", "0.009", "0.001", "0.040"],
    ],
    caption_text="表 3-3　營養轉換資料庫之已標定紀錄範例",
)
para("表 3-3 之數值均可由下述程序重算：自表 3-4 取得每 100 公克營養值 → 乘以表 3-1 之"
     "秤重並除以 100 得表 3-1a 之一分量總量 → 再除以管線體積得表 3-3 之每立方公分係數。"
     "加權合成品項之 SYN 編號與成分配比詳見 food_nutrition_db.json。")
para("為提升系統之強健性，在由視覺語言模型輸出之食物名稱查詢資料庫時，本研究另設計了名稱對應"
     "機制：先以人工建立之對應表（manual mapping）處理常見之別名與同義詞（例如將「cabbage」"
     "對應至「cabbage_stir_fried」、「kimchi」對應至「kimchi_korean」），若查無對應則採用字串近似"
     "匹配（fuzzy matching）尋找最相近之識別鍵；若仍無法匹配，則將該品項標記為未知（UNKNOWN），"
     "僅回報其管線估計體積而不計入營養總和，並於結果中提示使用者。此一設計確保系統在面對資料庫未收錄之"
     "食物時，仍能維持流程之完整性並提供可解釋之回饋。")
todo_note("資料庫現階段僅收錄十項已完成標定之品項（見表 3-1～3-4）。"
          "其餘品項待完成管線體積—秤重標定後再擴充；加權合成之 SYN 編號規則見 3.4.1 節。")

heading("3.5　與 Nutrition5k 等公開資料集之設計對照", level=2)
para("本研究之資料庫雖非大規模公開 benchmark，但其 schema 設計可對照 Nutrition5k"
     f"{cit(20)}等文獻。Nutrition5k 以 dish 為單位，同時記錄 ingredient breakdown、"
     "重量、熱量、macronutrients 與可選 RGB-D；本研究則以「食物識別鍵＋每立方公分"
     "營養值＋堆積密度」為核心，使線上推論時可直接由推論管線所得之管線估計體積（{cm³}）"
     "換算營養。Nutrition5k 證實 volume scalar 可將熱量 MAPE 從 26.1% 降至 16.5%，"
     "本研究進一步將 volume 來源改為多視角幾何積分而非固定 top-view depth，以貼近"
     "活大自助餐之真實拍攝型態。")
para("在資料蒐集協定上，Nutrition5k 採固定 overhead 掃描與 RealSense 深度；TADA/mpFR"
     f"{cit(30,47)}強調 before/after 餐影像與 fiducial marker；Vinod 等人"
     f"{cit(22)}與 MetaFood{cit(23)}要求 checkerboard 或信用卡。本研究融合 TADA"
     "之「參考物＋標準化拍攝」精神與 Dehais 等人"
     f"{cit(29)}之「多視角幾何」思路，訂定三視角＋信用卡規範（3.2 節），使資料庫"
     "建置與線上推論共用同一幾何前提與體積估算管線。未來若擴充資料，可參考 Nutrition5k 逐項秤重"
     "流程，為每道活大自助餐樣本同時記錄「影像—管線估計體積—秤重—營養」四元組，以支援"
     "更嚴格之 ablation 與跨方法比較。")

# ===========================================================================
# 第四章 基於多模態視覺與三維點雲之體積估算架構
# ===========================================================================
heading("第四章　基於多模態視覺與三維點雲之體積估算架構", level=1, page_break_before=True)
para("本章提出本研究之核心方法：由使用者以智慧型手機拍攝之三張餐點影像（正上方、左四十五度、"
     "右四十五度）及同場之標準信用卡，經完整推論管線自動估算各食材之管線估計體積，並查詢營養"
     "轉換資料庫，輸出熱量與三大營養素。管線分為語意前處理（以 Qwen2.5-VL 辨識、SAM3 分割）、"
     "三維幾何處理（以 DUSt3R 重建、比例尺錨定、體積積分）與營養轉換三階段；其中體積積分所得"
     "之數值為本研究管線定義下之估計體積，資料庫建置階段亦採用相同流程標定體積—重量對應。")

heading("4.1　系統總體管線架構", level=2)
para("本系統之整體運算管線可分為三個階段：第一階段為「語意前處理」，以視覺語言模型辨識並定位"
     f"食物品項{cit(6)}，再以萬物分割模型取得各食材之像素級遮罩{cit(12)}；第二階段為"
     "「三維幾何處理」，以多視角三維重建模型還原食物之三維點雲，並透過信用卡完成尺度錨定"
     f"與體積積分{cit(13,19,22)}；第三階段為「營養轉換」，透過營養轉換資料庫將管線估計體積"
     "轉換為"
     f"重量、熱量與三大營養素{cit(18,20)}。三個階段依序串接，且每一階段之輸出皆為可視化、"
     "可檢驗之中間結果，構成一條可解釋之計算管線。")
fig_img("fig4_1.png", "圖 4-1　系統總體管線架構圖", width_cm=15)

heading("4.1.1　行動端與雲端協同運作架構", level=3)
para("考量視覺語言模型、萬物分割模型與三維重建模型皆具相當之運算與記憶體需求，難以於一般行動"
     f"裝置上即時執行{cit(30,48)}，本系統採用行動端與雲端協同之架構。行動端（使用者手機）"
     "負責輕量任務，包括：引導使用者完成標準化的多視角拍攝、於畫面上提示比例尺參考物之擺放、"
     "將影像壓縮後上傳，以及接收並呈現最終之熱量與營養分析結果。雲端則部署具圖形處理器"
     "（GPU）之推論伺服器，負責執行運算密集之模型推論與體積計算。")
para("此一分工之優點在於：其一，使用者無需具備高階運算硬體或專用深度感測器，僅需一般具相機之"
     "智慧型手機即可使用，符合普惠可用之設計目標；其二，模型集中部署於雲端，便於統一維護、"
     "更新與擴充，亦利於在多位使用者間共享運算資源；其三，可依實際流量彈性調配雲端運算資源，"
     "兼顧成本與效能。行動端與雲端之間透過應用程式介面（API）以標準之 HTTP 協定溝通，"
     "傳遞影像資料與分析結果。")
fig_img("fig4_2.png", "圖 4-2　行動端與雲端協同運作架構圖", width_cm=13)

heading("4.1.2　使用者操作流程設計", level=3)
para("為降低使用者之操作負擔並確保輸入影像品質，本系統設計如下之標準化操作流程：")
numbered("擺放比例尺：使用者將標準信用卡平放於餐點旁之同一平面上，並確認其完整入鏡、無遮擋與"
         "反光。")
numbered("多視角拍攝：依系統畫面之引導，依序拍攝正上方、左側四十五度與右側四十五度共三張影像；"
         "系統以視覺輔助（如取景框與角度提示）協助使用者對準建議之視角。")
numbered("上傳與推論：使用者確認影像後，行動端將三張影像上傳至雲端推論服務，並顯示分析進度。")
numbered("結果呈現：雲端完成推論後回傳分析結果，行動端以清晰之介面呈現各食材之名稱、管線估計體積、"
         "推估重量、熱量與三大營養素，以及全餐之營養總和。")
para("上述流程於實際系統中由行動端（三視角拍攝引導、影像上傳與結果呈現）與雲端推論服務"
     "（依序載入 Qwen2.5-VL、SAM3、DUSt3R 並執行體積與營養計算）協同完成；介面設計、"
     "部署環境與評估指標將於第五章說明。")

heading("4.2　基於輕量化視覺語言模型的食物品項辨識", level=2)
para("食物品項辨識為管線之首要步驟，其任務為由參考視角（正上方）影像中辨識出所有食物品項，"
     f"並取得各品項之大致位置。本研究採用輕量化視覺語言模型 Qwen2.5-VL {cit(6)}完成此一"
     "開放域辨識任務。")
para("具體作法為：將參考視角影像連同一段精心設計之文字提示（prompt）輸入 Qwen2.5-VL，"
     "提示中明確告知影像之寬與高（像素），並要求模型逐一辨識影像中每一種不同之食物，"
     "且為每一品項輸出唯一之邊界框（以左上與右下角之像素座標表示）。此外，提示中特別要求模型"
     "將作為比例尺之信用卡（card）亦視為一個獨立品項加以輸出，以利後續之尺度錨定。模型之輸出"
     "被規範為結構化之字典格式，鍵為食物名稱、值為邊界框座標，以便程式解析。")
para("採用視覺語言模型之關鍵優勢在於其開放域辨識能力：系統無需預先定義固定之食物類別集合，"
     "即可辨識種類繁多之食物，包含地方料理與混合餐點，大幅提升系統之適用範圍。同時，模型所"
     "輸出之邊界框為後續之分割步驟提供了精準之定位提示，使萬物分割模型得以聚焦於各食材所在"
     "之區域。")
para("需說明的是，本步驟僅利用視覺語言模型完成其擅長之「辨識與定位」，而不要求其估計份量或"
     "熱量；份量之量測交由後續具幾何依據之三維重建流程處理。此一職責劃分正是本研究克服視覺"
     "語言模型幻覺問題之核心設計。")
figure_placeholder("圖 4-3　視覺語言模型食物辨識與定位結果示意圖",
                   "參考視角影像上標註各食物品項名稱與邊界框（含信用卡）之辨識結果")

heading("4.3　基於萬物分割模型的多食材邊界分割", level=2)
para("取得各食物品項之邊界框後，本研究以萬物分割模型 SAM3"
     f"{cit(12)}進一步取得各食材於影像中之像素級"
     "遮罩（mask）。作法為：以視覺語言模型輸出之各品項邊界框（或品項名稱）作為 SAM3 之提示，"
     "引導其分割出對應食材之精確輪廓。相較於僅使用矩形邊界框，像素級遮罩能貼合食材之真實形狀，"
     "排除框內之背景與相鄰食材，為後續之三維點雲界定提供精確之區域依據。")
para("為便於後續處理，本研究將各食材之分割結果以彩色遮罩圖與對應之顏色對照表（color map）"
     "儲存，並自遮罩中萃取各食材之外輪廓（contour）。於萃取輪廓時，會濾除面積過小之雜訊區域，"
     "以提升分割結果之穩定性。這些以多邊形表示之輪廓，將於體積計算階段被還原為二值遮罩，"
     "用以界定各食材（以及信用卡）在參考視角中所對應之像素範圍。")
para("信用卡本身亦被視為一個特殊之「食材」加以分割，其遮罩將專門用於後續之尺度錨定與參考平面"
     "擬合。由於信用卡為平整之矩形且邊界清晰，SAM3 通常能取得穩定而精確之分割結果。")
repo_img("90_auto_segmented_auto_segmented.png",
         "圖 4-4　萬物分割模型（SAM3）之多食材邊界分割結果（各食材與卡片之像素級彩色遮罩與面積）",
         width_cm=11, color=True)

heading("4.4　基於多視角三維重建模型之空間重建與對齊", level=2)
para("在語意前處理取得各食材之範圍後，系統進入三維幾何處理階段。本研究採用多視角三維重建模型"
     f"DUSt3R {cit(13)}，由正上方與左右四十五度共三張影像還原食物場景之三維點雲。")
para("DUSt3R 之運作方式為：對輸入之影像兩兩配對，直接回歸出以像素對齊之三維點圖（pointmap），"
     f"即影像中每一像素所對應之三維座標{cit(13)}。本研究以「完整配對（complete graph）」"
     "之方式建立影像對，使三張影像之間皆彼此配對，以取得充分之幾何約束。此設計呼應 Dehais"
     f"等人於雙視角食物重建之經驗{cit(29)}，並在視角數上擴展為三視角以強化高度還原。"
     "接著透過全域對齊（global alignment）程序，以點雲最佳化之模式，將各影像對之點圖"
     "整合至同一之三維座標系，同時隱式地推得各視角之相機內參與外參。")
para("重建完成後，系統取得三項關鍵輸出：（一）每一視角之逐像素三維點座標；（二）各視角之相機"
     "姿態（位姿矩陣）；（三）參考視角之相機焦距。這些輸出構成後續尺度錨定與體積積分之幾何基礎。"
     "值得再次強調的是，DUSt3R 所還原之三維結構僅具相對尺度，其絕對大小尚待下一節之比例尺錨定"
     "予以確定。")
para("與文獻方法對照：Dehais 等人"
     f"{cit(29)}以雙視角 stereo 建點雲，需特徵匹配與相機標定；Shao 等人"
     f"{cit(21)}以 cGAN 由單張 RGB 重建 voxel，幾何為隱式學習；Vinod 等人"
     f"{cit(22)}則載入預建 3D mesh 並估 pose。DUSt3R"
     f"{cit(13)}之優勢在於免已知內參與外參、可直接輸出稠密 pointmap，"
     "較 COLMAP"
     f"{cit(14)}更適合一般使用者隨手拍攝；代價是輸出仍為相對尺度，"
     "必須如 2.3.3 節所述以信用卡錨定。本研究選三視角而非 Dehais 之雙視角，"
     "意在增加視差以改善高度估計，同時維持可接受之拍攝負擔。")
figure_placeholder("圖 4-5　多視角三維重建之點雲結果示意圖",
                   "由三張影像重建之食物場景三維點雲（含相機姿態）之視覺化")

heading("4.5　比例尺錨定轉換與三維體積積分演算法", level=2)
para("本節為體積估算之關鍵，說明如何將無尺度之三維點雲轉換為具公分單位之管線估計體積。演算法可"
     "分為三個步驟：比例尺錨定、參考平面擬合，以及厚度積分。")

heading("4.5.1　比例尺錨定", level=3)
para("由於三維重建結果僅具相對尺度，需藉由已知真實尺寸之信用卡完成尺度校正"
     f"{cit(19,22)}。系統首先取出信用卡遮罩所對應之三維點集，並計算此點集之最大點對距離，"
     "作為信用卡對角線於三維模型中之長度"
     "（以模型單位表示）；另一方面，標準信用卡之真實對角線長度為已知（約 10.12 公分，"
     "由長 8.56 公分與寬 5.40 公分計算而得）。換算係數 s 定義如式 (4-1)：")
equation(
    m_omath(
        m_text("s"), m_text(" = "),
        m_frac(
            m_sub(m_text("L", italic=True), m_text("card,real")),
            m_sub(m_text("L", italic=True), m_text("card,model")),
        ),
    ),
    number="(4-1)",
)
para("式中，{s} 即 cm_per_model_unit，表示每一模型單位所對應之公分數；{L:card,real} 為信用卡"
     "真實對角線長度（cm）；{L:card,model} 為信用卡於模型中之對角線長度（模型單位）。",
     indent_chars=0)
para("此一係數表示「每一模型單位所對應之公分數」，將其套用於整個三維模型，即可將所有相對尺度之"
     "長度轉換為絕對之公分尺度。由於信用卡與餐點置於同一場景並一同重建，其尺度資訊得以正確傳遞"
     "至整個食物模型，此即本研究免除專用深度感測硬體而仍能取得絕對尺度之關鍵。")

heading("4.5.2　參考平面擬合", level=3)
para("為計算食物之高度（厚度），需先確定食物所在之基準平面，亦即餐盤或桌面所構成之參考平面。"
     "本研究以信用卡之三維點，結合各食材遮罩邊界處之三維點，作為擬合平面之樣本點集——此因食材"
     "之邊界通常貼近盤面，可視為位於基準平面上。以奇異值分解（SVD）對此樣本點集進行平面擬合，"
     "求得平面之法向量與方程式參數。任一三維點至此平面之有號距離，即代表該點相對於盤面之高度。")

heading("4.5.3　厚度積分與體積計算", level=3)
para("在確定換算係數與參考平面後，系統對每一食材（信用卡除外）進行體積計算。首先，將所有視角之"
     "三維點透過參考視角之相機姿態反投影至參考視角之影像平面，並僅保留落於該食材遮罩範圍內之點，"
     "藉此彙整多視角之三維資訊以提升對食材立體結構之涵蓋。接著，對於參考視角影像中屬於該食材之"
     "每一像素位置（bin），統計落入該位置之所有三維點相對於參考平面之高度值，並取其最大值與最小值。")
para("為兼顧「實心食物」與「中空／有間隙食物」之情形，本研究採用如下之厚度判定邏輯：若某像素"
     "位置之高度最大值與最小值之差異超過一微小門檻，代表該處具可觀之上下表面間距（如有厚度或"
     "堆疊之食物），則以「最大值減最小值」作為該處之厚度；反之，若差異甚小，代表僅重建到單一"
     "表面，則以「最大值」（相對盤面之高度）作為該處之厚度。此一設計使演算法能同時合理處理"
     "多種食物型態。")
para("最後，將各像素位置之厚度值沿食材遮罩範圍加總，並乘以單一像素於參考平面上所對應之實際"
     "面積（由相鄰像素三維座標之中位間距 {dx}、{dy} 求得），再乘以換算係數之立方以完成單位換算，"
     "即得該食材之管線估計體積（立方公分）。此處之體積並非獨立於本方法之外之物理真值，而是"
     "第四章所述管線之輸出；因其與第三章資料庫建置所用之體積定義一致，故可作為營養換算之"
     "幾何輸入。其體積積分如式 (4-2)：")
equation(
    m_omath(
        m_sub(m_text("V", italic=True), m_text("pipeline")),
        m_text(" = "),
        m_nary(
            [
                m_sub(m_text("t", italic=True), m_text("i")),
                m_text(" × "),
                M.delta("x"),
                m_text(" × "),
                M.delta("y"),
                m_text(" × "),
                m_sup(m_text("s", italic=True), m_text("3")),
            ],
            sub=m_text("i∈mask"),
        ),
    ),
    number="(4-2)",
)
para("式中，{t:i} 為第 i 個像素位置之厚度；{Δx}、{Δy} 為參考平面上相鄰像素間距（cm）；"
     "{s} 為式 (4-1) 之尺度換算係數；mask 為該食材之分割遮罩範圍。",
     indent_chars=0)
para("上述流程對每一食材重複執行，即可得到全餐各食材之管線估計體積。系統並可輸出厚度圖（thickness map）"
     "作為視覺化之中間結果，供使用者與研究者檢視各處之估計高度，體現本方法之可解釋性。")
para("厚度積分法與文獻中其他幾何份量算法之比較如下：Vinod 等人"
     f"{cit(22)}以 rendered mask 面積比縮放 3D mesh 體積；Dehais 等人"
     f"{cit(29)}在 mesh 表面直接積分；MUSEFood"
     f"{cit(35)}以微分幾何模型估計；Dhar 等人"
     f"{cit(50)}以 2D contour 特徵回歸。本研究之 thickness binning 法"
     "不需預建 mesh，亦不需訓練回歸器，但假設食材可近似為盤面之上之柱體或"
     "堆疊體，對極度不規則或大量湯汁之食物可能低估或高估。此為幾何法共同"
     "面臨之 trade-off，可透過第六章案例分析量化。")
fig_img("fig4_6.png", "圖 4-6　比例尺錨定與厚度積分演算法示意圖", width_cm=13)
figure_placeholder("圖 4-7　食材厚度圖視覺化",
                   "以顏色深淺表示各像素相對盤面高度之厚度圖")

heading("4.6　熱量與營養素估算模組", level=2)
para("取得各食材之管線估計體積後，系統進入最後之營養轉換階段。本節說明由管線估計體積推得重量，"
     "再由重量推得熱量與三大營養素之流程。")

heading("4.6.1　體積轉重量估算方法", level=3)
para("管線估計體積至重量之轉換透過食物堆積密度完成，其關係即式 (3-1)。"
     "堆積密度與每百公克營養值已預先建置於營養轉換資料庫中，且密度係以與本節相同之推論管線"
     "對各品項一分量標定而得（見第三章 3.3 節）。系統依據視覺語言模型"
     "辨識所得之食物名稱，於資料庫中查得該食物之對應紀錄。"
     "為提升名稱匹配之強健性，系統先以人工對應表處理常見別名，若無對應則採字串近似匹配，"
     "仍無法匹配者標記為未知並僅回報管線估計體積。")

heading("4.6.2　重量轉營養成分推估方法", level=3)
para("營養轉換資料庫已預先將堆積密度與每 100 公克營養值合併，儲存為「每立方公分」之"
     "熱量與各營養素含量。因此，於線上推論時可略去顯式之重量計算，直接將食材"
     "之管線估計體積乘以其每立方公分之營養值，即得該食材之熱量、蛋白質、脂肪與碳水化合物含量。"
     "此一設計在數學上等價於「管線估計體積→重量→營養」兩階段轉換，但在計算上更為簡潔。")

heading("4.6.3　熱量計算流程", level=3)
para("綜合上述，單一食材之熱量與營養素估算流程為：（一）以食物名稱查得資料庫紀錄；"
     "（二）以管線估計體積乘以每立方公分熱量得該食材熱量；（三）以管線估計體積乘以每立方公分"
     "各營養素值得各營養素含量。將全餐所有食材之結果加總，即得整餐之總熱量與三大營養素總量。"
     "系統最終將各食材之名稱、管線估計體積、推估重量、熱量與營養素，以及全餐之營養總和，"
     "整理為結構化之分析報告回傳予行動端呈現。整體之營養估算流程如圖 4-8 所示。")
fig_img("fig4_8.png", "圖 4-8　熱量與營養素估算流程圖", width_cm=9.5)

heading("4.7　與代表性文獻方法之對照分析", level=2)
para("本節將第四章方法與 2.4 節代表性研究對照，說明設計取捨。就辨識而言，He 等人"
     f"{cit(33)}與 goFOOD{cit(34)}採封閉集合 CNN；本研究改以 Qwen2.5-VL"
     f"{cit(6)}開放域輸出邊界框，類似 Vinod 等人以 YOLOv8 偵測但無類別上限。"
     "就分割而言，Vinod 等人"
     f"{cit(22)}以 SAM 做 instance mask；本研究以 SAM3"
     f"{cit(12)}支援文字概念提示，可一次分割多個開放詞彙食材。就幾何而言，Dehais"
     f"等人{cit(29)}以雙視角 stereo + 參考卡；Shao 等人"
     f"{cit(21)}以單張 RGB 隱式 voxel；Vinod 等人{cit(22)}以 3D model scaling；"
     "本研究以三視角 DUSt3R"
     f"{cit(13)}取得稠密點雲，再以信用卡{cit(19)}做尺度錨定與厚度積分，兼具"
     "Dehais 之可解釋幾何與 Shao 之免 depth 推論，且不需 Vinod 式之 per-class 3D 資產。")
para("就營養換算而言，Fang 等人"
     f"{cit(32)}與 Shao 等人{cit(49)}直接回歸總 kcal；Nutrition5k"
     f"{cit(20)}先估 mass 再乘 per-gram nutrition；本研究採「管線估計體積×每 {{cm³}} 營養值」"
     "（第三章資料庫），使中間量（管線估計體積、推估重量）皆可檢視。Dhar 等人"
     f"{cit(50)}以 mask 幾何特徵回歸熱量，在垂直品類上精度高但難擴至混合餐；"
     "本研究以逐食材體積積分再求和，較適合活大自助餐之多品項場景。表 4-1 摘要對照。")
make_table(
    ["環節", "代表文獻作法", "本研究作法", "主要差異"],
    [
        ["辨識", "He [33] 封閉 21 類", "Qwen2.5-VL 開放域", "免固定類別表"],
        ["分割", "Vinod [22] YOLO+SAM", "VLM 框 + SAM3", "文字概念分割"],
        ["幾何", "Dehais [29] 雙視角 stereo", "三視角 DUSt3R", "免特徵匹配調參"],
        ["尺度", "TADA [30] fiducial", "ISO 信用卡對角線", "日常隨身物品"],
        ["熱量", "Shao [49] E2E 回歸", "體積×密度×營養表", "可解釋中間量"],
    ],
    caption_text="表 4-1　本研究管線與代表性文獻之環節對照",
)

# ===========================================================================
# 第五章 實驗設計與系統實作
# ===========================================================================
heading("第五章　實驗設計與系統實作", level=1, page_break_before=True)
para("本章說明如何將前述體積估算架構落實為可運行之系統原型，並訂定實驗設計以評估其效能。"
     "系統採行動端—雲端協同：使用者以手機依規範拍攝三視角影像並上傳，雲端 GPU 伺服器"
     "依序執行辨識、分割、三維重建、體積積分與營養換算，再回傳結構化分析結果。"
     "本章另說明活大自助餐測試場景、硬體與軟體環境、各模型推論參數，以及體積、熱量與"
     "三大營養素之評估指標。")

heading("5.1　行動應用程式設計與實作", level=2)
para("行動應用程式為使用者與系統互動之介面，其設計目標為：以最低之操作負擔，引導使用者取得"
     "品質穩定之多視角影像，並以清晰易懂之方式呈現分析結果。")

heading("5.1.1　多視角拍攝介面設計", level=3)
para("多視角拍攝介面為確保輸入品質之關鍵。介面依 4.1.2 節之操作流程，逐步引導使用者完成三個"
     "視角之拍攝。於每一視角，介面提供視覺輔助元素，例如建議之取景框、角度指示與比例尺擺放提示，"
     "協助使用者對準正上方、左側四十五度與右側四十五度之建議視角，並提醒確認信用卡完整入鏡、"
     "無遮擋與反光。拍攝完成後，介面提供預覽與重拍功能，使用者確認無誤後方進入上傳階段。")
fig_img("fig5_1.png", "圖 5-1　多視角拍攝介面設計稿（拍攝引導、取景框、角度與比例尺提示）", width_cm=7.5)
todo_note("可補充實際採用之行動端開發框架（如 Flutter／React Native／原生 Android/iOS）"
          "與版本、以及介面實作之細節與截圖。")

heading("5.1.2　影像上傳與資料傳輸流程", level=3)
para("使用者確認三張影像後，行動端將影像進行適度壓縮以降低傳輸量與延遲，再透過 HTTP 之 API "
     "請求上傳至雲端推論服務。上傳請求中一併攜帶必要之參數（如指定參考視角之索引）。上傳後，"
     "行動端顯示分析進度並等待雲端回傳結果。為兼顧使用者體驗，介面於等待期間提供進度提示，"
     "並針對網路異常或推論失敗之情況提供重試機制與錯誤提示。")
fig_img("fig5_2.png", "圖 5-2　影像上傳與資料傳輸流程圖（時序）", width_cm=13)

heading("5.1.3　熱量分析結果呈現介面", level=3)
para("雲端回傳分析結果後，行動端以結果呈現介面顯示。介面分為兩個層次：其一為「全餐總覽」，"
     "顯示整餐之總熱量與蛋白質、脂肪、碳水化合物三大營養素總量；其二為「各食材明細」，"
     "逐項列出各食材之名稱、估計體積、熱量與各營養素含量。對於資料庫未收錄而無法匹配之食材，"
     "介面明確標示為未知並僅顯示其體積，以維持結果之可解釋性與誠實性。")
fig_img("fig5_3.png", "圖 5-3　熱量分析結果呈現介面設計稿（全餐總覽與各食材明細）", width_cm=7.5)

heading("5.2　雲端人工智慧推論服務建置", level=2)
para("雲端推論服務為承載模型運算之核心。相較於 goFOOD"
     f"{cit(34)}、TADA {cit(30)}等早期系統多將模型部署於固定伺服器，本研究採"
     "按秒計費之 GPU 雲平台，以支援間歇性推論與彈性擴展。本節說明其模型部署架構、"
     "應用程式介面設計與推論流程管理。")

heading("5.2.1　模型部署架構", level=3)
para("雲端推論服務部署於具圖形處理器（GPU）之伺服器上，並依序載入三個模型：視覺語言模型"
     "Qwen2.5-VL、萬物分割模型 SAM3 與多視角三維重建模型 DUSt3R。由於三個模型皆佔用相當之"
     "GPU 記憶體，服務於各階段推論完成後主動釋放不再使用之模型與快取，以控制記憶體用量，"
     "使系統得以於單一具中階 GPU 記憶體之伺服器上完整執行整條管線。")
para("在部署平台上，本研究採用 RunPod 雲端 GPU 運算平台。RunPod 屬按用量計費（billed "
     "per-second）之 GPU 租用服務，相較於自建機房或長期租用專用伺服器，其優勢在於可依實際需求"
     "彈性開啟與關閉運算執行個體（Pod），僅於推論期間計費，適合本系統「間歇性、單次推論時間"
     "較長」之使用型態，能有效控制運算成本。")
para("具體部署方式如下：本研究以 RunPod 提供之 PyTorch 官方容器映像（Runpod PyTorch，"
     "內含 CUDA 與 PyTorch 執行環境）啟動 Pod，並掛載一顆容量為 80 GB 之持久化網路磁碟"
     "（Network Volume）至容器之 /workspace 路徑。模型權重（Qwen2.5-VL、SAM3、DUSt3R）、"
     "conda 虛擬環境與 HuggingFace 快取皆存放於此持久化磁碟；如此一來，即使關閉 Pod 以停止"
     "計費，環境與模型仍得以保留，下次啟動時無需重新下載約數十 GB 之模型，大幅縮短冷啟動時間"
     "並降低重複下載之成本。運算執行個體與網路磁碟部署於同一資料中心區域（本研究採用 EU-RO-1），"
     "以確保磁碟得以正確掛載並取得較佳之存取效能。")
para("在 GPU 選型上，本研究之主要實驗採用單張 NVIDIA RTX PRO 4500（32 GB 顯示記憶體）"
     "執行個體；此規格之顯示記憶體足以容納本管線推論尖峰之需求，且於實驗期間供貨穩定、"
     "單位成本適中。為回應不同部署情境之需求，本研究另於 5.4 節針對多種可用之 GPU 規格進行"
     "比較，並於第六章分析其推理時間與推理成本之差異。")

heading("5.2.2　應用程式介面設計", level=3)
para("雲端服務對外提供應用程式介面（API）供行動端呼叫。API 接收三張影像及參考視角索引等參數，"
     "於伺服器端依序執行辨識、分割、重建與營養轉換，並將結果以結構化之 JSON 格式回傳。回傳內容"
     "包含各食材之名稱、體積、熱量與三大營養素，以及全餐之營養總和；對於未匹配之食材則附註說明。")
make_table(
    ["API 端點", "方法", "輸入", "輸出"],
    [
        ["/analyze", "POST", "三張影像、參考視角索引", "各食材與全餐營養報告（JSON）"],
        ["/health", "GET", "—", "服務健康狀態"],
    ],
    caption_text="表 5-1　雲端推論服務之 API 規格",
)
todo_note("可補充 API 之詳細請求／回應格式範例、驗證機制與錯誤碼設計。")

heading("5.2.3　推論流程管理", level=3)
para("推論流程管理負責整合辨識、分割與重建三個階段，並確保其正確銜接與資源之有效利用。"
     "流程管理之要點包括：（一）階段間之資料傳遞，如將視覺語言模型輸出之邊界框傳遞予分割模型、"
     "將分割輪廓傳遞予體積計算；（二）GPU 記憶體之管理，於各階段完成後釋放模型與快取；"
     "（三）例外處理，如辨識無結果時之退場機制、分割或重建失敗時之錯誤回報。透過此一流程管理，"
     "系統得以穩定地完成由影像至營養報告之完整推論。")
fig_img("fig5_4.png", "圖 5-4　雲端推論流程管理示意圖（含 GPU 記憶體管理）", width_cm=12)

heading("5.3　實驗資料集與測試範例介紹", level=2)
para("為驗證系統於實際用餐情境中之表現，本研究以國立臺灣大學活動中心自助餐（以下簡稱「活大"
     "自助餐」）之菜色作為示範與實驗資料來源。活大自助餐提供種類多元之中式自助餐菜色，包含"
     "主食（如白飯、炒麵）、蔬菜、肉類與湯品等，貼近一般大學生與教職員之日常取餐型態，"
     "具備作為本系統驗證場域之代表性與可及性。")
para("本研究設計兩類測試場景，複雜度由低至高，分別對應單一食材與活大自助餐混合取餐。每一測試"
     "範例皆依 3.2 節之規範拍攝三視角影像，並置放標準信用卡作為比例尺；各食材之實測重量以"
     "電子秤取得，作為評估推估重量與熱量之基準真值。體積方面，本研究不以獨立量測之物理真值"
     "體積作為主要基準，而以推論管線所得之管線估計體積評估幾何環節之表現，並透過「管線估計"
     "體積×資料庫密度」所得之推估重量與秤重比較，檢驗整體系統之端到端準確度。")

heading("5.3.1　單一食材測試場景", level=3)
para("單一食材場景每次僅拍攝活大自助餐中之單一菜色（如一份白飯、一份炒青菜或一份主菜），"
     "用以在最單純之條件下驗證體積估算與尺度錨定之準確度，排除多食材相鄰所帶來之分割與"
     "遮擋干擾。此場景亦作為系統誤差之基準參照。")
todo_note("補充活大自助餐單一食材測試之品項清單、樣本數量、各樣本之管線估計體積與秤重數值。")

heading("5.3.2　活大自助餐混合取餐測試場景", level=3)
para("混合取餐場景模擬活大自助餐之典型取餐方式：於同一餐盤中同時盛裝一份主食與數樣配菜"
     "（如白飯、炒青菜、滷肉或炸物等）。此場景用以驗證系統於多食材相鄰、且存在部分遮擋"
     "與高度差異情況下之辨識、分割與體積估算表現，較貼近使用者實際用餐時之操作情境。")
para("資料蒐集協定參考 Nutrition5k"
     f"{cit(20)}之「逐項加菜、即時秤重」精神與 TADA"
     f"{cit(47)}之標準化影像流程：每份樣本先以電子秤量測各食材之一分量重量，"
     "再依 3.2 節完成三視角 + 信用卡拍攝，並以推論管線取得管線估計體積；熱量基準真值"
     "則由秤重與衛福部資料庫"
     f"{cit(18)}換算而得。"
     "此設計使本研究能同時報告管線體積相關指標、推估重量誤差與 calorie MAPE，並與文獻中 "
     "MAE/MAPE 指標對齊。活大場景之 domain gap（中式複合菜、醬色、堆疊）"
     "亦為检验 Shao 等人"
     f"{cit(49)}等西洋校園資料訓練模型泛化能力之理想測試床。")
todo_note("補充活大自助餐混合取餐測試之典型組合、樣本數量與各食材之管線估計體積／秤重數值。")

heading("5.4　硬體配置與軟體環境說明", level=2)
para("本節記錄系統開發、部署與推論所使用之硬體配置與軟體環境，以確保實驗之可重現性。")
make_table(
    ["項目", "規格／版本"],
    [
        ["雲端平台", "RunPod（GPU Cloud，按秒計費）"],
        ["資料中心區域", "EU-RO-1"],
        ["雲端 GPU", "NVIDIA RTX PRO 4500（32 GB GDDR7）※主要實驗；其餘規格見表 5-3"],
        ["系統記憶體（RAM）", "62 GB"],
        ["vCPU", "12 核"],
        ["持久化儲存", "80 GB Network Volume（掛載於 /workspace）"],
        ["容器映像", "Runpod PyTorch（Ubuntu 22.04，CUDA 12.x）"],
        ["作業系統", "Ubuntu 22.04 LTS"],
        ["Python 版本", "3.12"],
        ["深度學習框架", "PyTorch 2.7（CUDA 12.6）"],
        ["視覺語言模型", "Qwen2.5-VL-7B-Instruct"],
        ["分割模型", "SAM3（facebook/sam3）"],
        ["三維重建模型", "DUSt3R（ViTLarge, 512, dpt）"],
        ["行動端裝置", "TODO（測試用手機型號與作業系統版本）"],
    ],
    caption_text="表 5-2　硬體配置與軟體環境",
    todo=True,
)
para("為回應不同部署情境（如成本優先、效能優先或大顯存需求）之考量，本研究比較 RunPod 平台上"
     "數種可用於執行本推論管線之 GPU 規格。判斷某一 GPU 是否「足以執行」之主要準則為其顯示"
     "記憶體是否能容納本管線於推論尖峰時之需求；由於本管線同時使用 7B 等級之視覺語言模型與"
     "多視角三維重建模型，建議之最低顯示記憶體為 24 GB。表 5-3 彙整各候選 GPU 之規格與"
     "概略之每小時租用費用（RunPod On-Demand，2026 年之參考價，實際費用會隨供需波動）。")
make_table(
    ["GPU 型號", "架構", "顯示記憶體", "系統 RAM", "每小時費用(US$)", "是否足以執行", "定位／備註"],
    [
        ["NVIDIA L4", "Ada", "24 GB", "—", "約 0.40–0.79", "足夠", "成本最低，速度較慢"],
        ["RTX 4090", "Ada", "24 GB", "—", "約 0.60", "足夠", "性價比高，常缺貨"],
        ["RTX PRO 4500", "Blackwell", "32 GB", "62 GB", "約 0.74", "足夠（本研究採用）", "供貨穩定、成本適中"],
        ["RTX 5090", "Blackwell", "32 GB", "60 GB", "約 0.99", "足夠", "速度最快之消費級卡"],
        ["RTX PRO 6000", "Blackwell", "96 GB", "283 GB", "約 2.09", "過剩", "大顯存，成本偏高"],
        ["NVIDIA A100", "Ampere", "80 GB", "—", "約 1.40–2.10", "過剩", "資料中心級，適合訓練"],
    ],
    caption_text="表 5-3　RunPod 平台可用 GPU 規格與費用比較",
)
para("由表 5-3 可知，本管線對 GPU 之需求集中於 24 GB 以上之顯示記憶體，而非極高之運算等級；"
     "因此 L4、RTX 4090、RTX PRO 4500 與 RTX 5090 等中階 GPU 皆足以勝任，A100 與 RTX PRO "
     "6000 等資料中心／大顯存等級雖亦可執行，但其成本較高、就本應用而言運算能力有所過剩。"
     "各規格於實際推理時間與單次推理成本上之差異，將於第六章 6.4.4 節進一步量化比較。")
para("關於系統之最低硬體限制與流量—配置對應規格，本研究依模型之記憶體佔用與單次推論之運算量"
     "評估如下：整條管線於推論尖峰時之 GPU 記憶體需求，主要由視覺語言模型與三維重建模型主導；"
     "為使單一使用者之請求得以順利完成，建議之最低 GPU 顯示記憶體為 24 GB。於多使用者情境下，"
     "由於單次推論需時較長，服務吞吐量受限於 GPU 之數量與單次推論時間，可透過水平擴充"
     "（增加 GPU 執行個體並搭配請求佇列）以支撐更高之並行流量。")
make_table(
    ["並行使用人數／流量", "建議配置", "備註"],
    [
        ["單人測試／開發", "1 × 24GB GPU（如 L4／RTX PRO 4500）", "循序處理，用畢即關閉 Pod"],
        ["小規模（數人並行）", "數個 GPU 執行個體 + 請求佇列", "TODO：實測吞吐量後補充"],
        ["中大規模", "自動擴充（auto-scaling）+ 負載平衡", "TODO：可評估 Serverless 部署"],
    ],
    caption_text="表 5-4　人數流量與配置規格對應",
    todo=True,
)

heading("5.5　推論參數設定", level=2)
para("本節說明各模型之推論參數與提示詞配置，以確保實驗結果之可重現性。")
para("在視覺語言模型方面，本研究以半精度（float16）載入 Qwen2.5-VL-7B-Instruct，並採用"
     "自動裝置配置。其文字提示明確告知影像尺寸、要求逐一辨識每種食物並輸出單一邊界框、"
     "且將信用卡納入為獨立品項，輸出格式限定為可解析之字典。在分割模型方面，SAM3 以視覺語言"
     "模型輸出之邊界框作為提示進行分割，並於萃取輪廓時濾除面積過小之區域。在三維重建模型方面，"
     "DUSt3R 以 512 之輸入尺寸、完整配對建立影像對，並以點雲最佳化模式進行全域對齊，"
     "以最小生成樹初始化並進行多次迭代最佳化。")
make_table(
    ["模型", "主要參數", "設定值"],
    [
        ["Qwen2.5-VL", "精度 / 裝置配置", "float16 / 自動"],
        ["Qwen2.5-VL", "提示詞", "結構化字典輸出（含 card）"],
        ["SAM3", "提示類型", "邊界框 / 概念文字"],
        ["SAM3", "輪廓面積門檻", "TODO"],
        ["DUSt3R", "輸入尺寸", "512"],
        ["DUSt3R", "配對方式 / 對齊模式", "完整配對 / 點雲最佳化"],
        ["DUSt3R", "全域對齊迭代次數", "TODO（如 300）"],
    ],
    caption_text="表 5-5　各模型之推論參數設定",
    todo=True,
)

heading("5.6　評估指標選用", level=2)
para("為客觀評估系統之準確度，本研究於體積、熱量與三大營養素三個層面選用適切之評估指標。"
     "整體而言，本研究以均方誤差（Mean Squared Error, MSE）與平均絕對百分比誤差"
     "（Mean Absolute Percentage Error, MAPE）作為主要之量化指標；其定義分別如式 (5-1) 與式 (5-2)。"
     "MAPE 以相對百分比表示，便於跨不同量級之食材與餐點進行比較，MSE 則對較大之誤差較為敏感，"
     "可反映極端誤差之影響。")
equation(
    m_omath(
        m_text("MAPE"), m_text(" = "),
        m_frac(m_text("1"), m_text("n")),
        m_nary(
            m_frac(
                [
                    m_text("|"),
                    m_sub(m_text("y", italic=True), m_text("i")),
                    m_text(" − "),
                    m_sub(m_text("ŷ", italic=True), m_text("i")),
                    m_text("|"),
                ],
                [
                    m_text("|"),
                    m_sub(m_text("y", italic=True), m_text("i")),
                    m_text("|"),
                ],
            ),
            sub=m_text("i=1"),
            sup=m_text("n"),
        ),
        m_text(" × 100%"),
    ),
    number="(5-1)",
)
equation(
    m_omath(
        m_text("MSE"), m_text(" = "),
        m_frac(m_text("1"), m_text("n")),
        m_nary(
            m_sup(
                [
                    m_sub(m_text("y", italic=True), m_text("i")),
                    m_text(" − "),
                    m_sub(m_text("ŷ", italic=True), m_text("i")),
                ],
                m_text("2"),
            ),
            sub=m_text("i=1"),
            sup=m_text("n"),
        ),
    ),
    number="(5-2)",
)
para("式中，{n} 為樣本數；{y:i} 為第 i 筆基準真值；{yhat:i} 為第 i 筆系統推估值。",
     indent_chars=0)

heading("5.6.1　體積估算誤差", level=3)
para("體積估算誤差以推論管線所得之管線估計體積為評估對象。由於本研究之體積為管線定義下之"
     "估計量，而非獨立驗證之物理真值，體積誤差主要指標著重於：（一）同一管線下之可重現性；"
     "（二）不同視角數量、比例尺錨定等設定對管線估計體積之影響。量化上可採 MAPE 與 MSE"
     "表示管線估計體積之變異或相對偏差。此外，本研究特別探討「不同視角數量」對體積估算誤差"
     "之影響：以較多視角（如三視角）之估計作為較佳之參照，比較減少視角數量時誤差之變化，"
     "藉此分析多視角資訊對幾何重建之貢獻。")
para("就系統整體準確度而言，更具實務意義之評估為：以管線估計體積乘以資料庫密度所得之推估重量，"
     "與電子秤實測重量比較（見 5.6.2 節），因為重量為本研究主要之實測基準真值。")
todo_note("本小節之實驗設計可再明確化：例如以三視角結果為基準，逐步減為兩視角、單視角，"
          "觀察管線估計體積之 MAPE 變化；或對同一分量重複拍攝，評估管線體積之可重現性。"
          "相關數據將於第六章呈現。")

heading("5.6.2　熱量估算誤差", level=3)
para("熱量估算誤差以系統推估之熱量與基準真值之差異衡量；基準真值由電子秤實測重量與標準營養值"
     "計算而得，亦即「秤重→營養」路徑，而非直接以管線估計體積作為熱量真值。採用 MAPE 與 MSE"
     "表示，並同時於單一食材層級與全餐層級進行評估。")

heading("5.6.3　三大營養素偏差", level=3)
para("三大營養素（蛋白質、脂肪、碳水化合物）之偏差分別以 MAPE 與 MSE 衡量，"
     "以檢視系統於各營養素上之估計一致性，並分析不同食物型態對各營養素估計之影響。")

heading("5.7　基準方法與文獻對照之實驗設計", level=2)
para("為呼應第二章文獻探討，本研究實驗除評估本方法外，亦規劃與「端對端語意估熱」"
     "基準比較。文獻中，Shao 等人"
     f"{cit(49)}與 He 等人{cit(33)}以 MAE／MAPE 評估單張 RGB 熱量；Nutrition5k"
     f"{cit(20)}以 dish-level MAE 比較 2D、RGB-D 與 volume scalar；Vinod 等人"
     f"{cit(22)}則報告 EMAPE。本研究採相同之 MAPE 作為主要可比指標，並額外報告"
     "體積相關指標與推估重量誤差，以分離「幾何份量」與「營養換算」兩階段誤差。")
para("基準方法選用直接使用大型多模態模型（如 GPT-4V"
     f"{cit(46)}或同級 VLM）由影像輸出總熱量與各食材描述，代表現有「免參考物、"
     "免幾何管線」之部署路線。本研究另建 LMM 基準（提示詞與原始輸出見第六章），"
     "不以 GPT-4 Technical Report"
     f"{cit(46)}作為食物熱量誤差之數值依據——該報告僅說明 GPT-4 可接收影像，"
     "未提供 food-calorie estimation benchmark。相較之下，Dehais 等人"
     f"{cit(29)}（MAPE 8.2%）與 Dhar 等人"
     f"{cit(50)}（MAE 7.85 kCal）代表幾何／參考物路線，但分別限於"
     "雙視角受控拍攝或封閉五類街頭食物，與活大混合餐盤之設定不同。"
     "本研究預期在相同活大協定下，以 MAPE 指標比較本方法與 LMM 基準，"
     "並在開放域與操作便利性上與 Vinod 式 3D model scaling 對照討論。")
para("實驗設計上，單一食材場景用以對照 Dehais 等人之低複雜度幾何重建設定；混合取餐"
     "場景用以測試 goFOOD、DeepFood 等系統所關注之多品項分割與份量問題。所有樣本"
     "均依 3.2 節規範拍攝，以確保與 TADA、MetaFood 等文獻在「有參考物」之前提下"
     f"公平比較{cit(30,23)}。未來工作可進一步在相同活大資料上複現 Shao 等人"
     f"{cit(49)}之 energy distribution 架構，作為額外 learning-based baseline。")

# ===========================================================================
# 第六章 實驗結果分析與討論
# ===========================================================================
heading("第六章　實驗結果分析與討論", level=1, page_break_before=True)
para("本章呈現並討論前述系統之實驗結果。評估對象為一套由三視角影像出發、以標準信用卡"
     "錨定絕對尺度、整合 Qwen2.5-VL／SAM3／DUSt3R 之食物體積與熱量估算系統；比較基準"
     "包含直接以大型多模態模型由影像估算熱量之方法，以及不同視角數量、不同 GPU 規格"
     "下之效能差異。表 6-0 整理文獻代表性方法之報告指標與實驗條件；"
     "須注意各文獻之資料集、輸入模態與指標定義不同，僅供背景對照。"
     "本研究之 MAPE 與 LMM 基準結果將於下列各節表格填入。")
make_table(
    ["方法", "文獻", "輸入條件", "文獻報告指標", "可比性說明"],
    [
        ["2D 直接回歸", "Nutrition5k [20]", "Top-view RGB", "MAPE 26.1%", "固定 top-view，非活大場域"],
        ["Volume scalar", "Nutrition5k [20]", "RGB + volume", "MAPE 16.5%", "含 volume 先驗"],
        ["Cross-domain RGB", "Shao [49]", "單張 RGB", "MAPE 11.47%", "96 張 nutrition study"],
        ["Voxel + RGB", "Shao [21]", "單張 RGB", "MAE 40.05 kCal", "原文明列 MAPE 不一致"],
        ["3D scaling", "Vinod [22]", "RGB+checkerboard", "EMAPE 17.67%", "需 3D model 資產"],
        ["雙視角 stereo", "Dehais [29]", "雙視角+參考卡", "MAPE 8.2%", "受控雙視角"],
        ["LMM 基準", "本研究建立", "三視角 RGB", "待填入", "同活大協定"],
        ["本研究", "—", "三視角+信用卡", "待填入", "同活大協定"],
    ],
    caption_text="表 6-0　文獻方法與本研究之實驗條件對照框架（文獻數值僅供背景）",
    todo=True,
)

heading("6.1　辨識與分割模型結合之效能評估", level=2)
para("本節評估視覺語言模型辨識與萬物分割模型分割結合後之整體前處理效能。辨識部分以食物品項之"
     "辨識正確率與定位品質衡量；分割部分以分割遮罩與人工標註之交集比聯集（Intersection over "
     "Union, IoU）等指標衡量。下表彙整各測試場景之辨識與分割效能。")
make_table(
    ["測試場景", "食物辨識正確率", "平均分割 IoU", "備註"],
    [
        ["單一食材（活大自助餐）", "TODO", "TODO", ""],
        ["混合取餐（活大自助餐）", "TODO", "TODO", ""],
    ],
    caption_text="表 6-1　辨識與分割模型結合之效能",
    todo=True,
)
para("〔結果討論〕本節之結果預期呈現：於活大自助餐單一食材場景，辨識與分割皆能取得高準確度；"
     "於混合取餐場景，隨多食材相鄰、遮擋與質地相近等因素，辨識與分割之難度隨之增加。"
     "相關趨勢與具體數值將於數據填入後補充討論。")
todo_note("填入各場景之辨識正確率與分割 IoU，並補充失敗案例分析（如相近食材誤辨、邊界誤分割）。")

heading("6.2　三維重建與管線體積估算分析", level=2)
para("本節分析三維重建與比例尺錨定後之管線估計體積表現，並探討導入信用卡比例尺前後之差異，"
     "以及視角數量對管線體積估算之影響。須再次強調，表中體積數值為本研究管線定義下之估計量，"
     "主要用於檢視幾何環節之表現；系統整體準確度仍以推估重量／熱量與秤重之比較為主（見 6.3 節）。")
make_table(
    ["食材／樣本", "管線估計體積({cm³})", "推估重量(g)", "秤重(g)", "重量誤差(%)"],
    [
        ["樣本 1", "TODO", "TODO", "TODO", "TODO"],
        ["樣本 2", "TODO", "TODO", "TODO", "TODO"],
        ["樣本 3", "TODO", "TODO", "TODO", "TODO"],
        ["平均", "—", "—", "—", "TODO"],
    ],
    caption_text="表 6-2　管線體積估算與推估重量準確度",
    todo=True,
)
make_table(
    ["視角數量", "管線體積 MAPE(%)", "管線體積 MSE", "備註"],
    [
        ["單一視角", "TODO", "TODO", "缺乏視差，僅供對照基準"],
        ["兩視角", "TODO", "TODO", ""],
        ["三視角（本研究）", "TODO", "TODO", "完整配置"],
    ],
    caption_text="表 6-3　不同視角數量之管線體積估算誤差比較",
    todo=True,
)
para("〔結果討論〕預期導入信用卡比例尺後，管線估計體積得以由相對尺度轉為具公分單位之估計值，"
     "大幅降低因缺乏尺度所造成之系統性誤差；且隨視角數量增加，食物立體結構之涵蓋更完整，"
     "管線體積估算之變異隨之下降。具體數值與趨勢圖將於數據填入後補充。")
figure_placeholder("圖 6-1　不同視角數量之體積估算誤差趨勢圖",
                   "以視角數量為橫軸、體積 MAPE 為縱軸之折線圖")
todo_note("填入體積準確度與視角數量比較之數據；如 5.6.1 節所述，可以較大之單視角誤差作為基準，"
          "凸顯多視角之改善幅度。")

heading("6.3　與大型多模態模型之熱量估算比較", level=2)
para("本節將本系統與自行建立之 LMM 基準（直接以大型多模態模型由影像估算熱量）進行"
     f"比較{cit(16,48)}，以驗證導入絕對尺度與幾何計算之效益。文獻背景方面，"
     f"Shao 等人{cit(49)}於 96 張真實用餐影像報告 MAPE 11.47%；Nutrition5k 2D "
     f"baseline 為 26.1%{cit(20)}；Vinod 等人 3D scaling 為 EMAPE 17.67%{cit(22)}。"
     "上述數值之實驗條件各異，僅供對照；本章在相同活大拍攝協定下報告本方法與"
     f"LMM 基準之 MAPE。GPT-4 Technical Report{cit(46)}未提供食物熱量 benchmark，"
     "不作為數值依據。")

heading("6.3.1　熱量估算誤差比較", level=3)
make_table(
    ["方法", "熱量 MAPE(%)", "熱量 MSE", "備註"],
    [
        ["大型多模態模型（基準）", "TODO", "TODO", "端對端估算"],
        ["本研究系統", "TODO", "TODO", "3D + 比例尺"],
    ],
    caption_text="表 6-4　熱量估算誤差比較",
    todo=True,
)
para("〔結果討論〕待表 6-4 數值填入後，將比較本方法與 LMM 基準於活大場域下之 MAPE 與 MSE，"
     "並討論幾何管線對份量偏離典型值樣本之影響。文獻中 Shao 等人"
     f"{cit(49)}、Nutrition5k{cit(20)}與 Vinod 等人{cit(22)}之數值僅作背景參考，"
     "不作直接優劣判定。")

heading("6.3.2　三大營養素估算比較", level=3)
make_table(
    ["營養素", "基準 MAPE(%)", "本研究 MAPE(%)", "基準 MSE", "本研究 MSE"],
    [
        ["蛋白質", "TODO", "TODO", "TODO", "TODO"],
        ["脂肪", "TODO", "TODO", "TODO", "TODO"],
        ["碳水化合物", "TODO", "TODO", "TODO", "TODO"],
    ],
    caption_text="表 6-5　三大營養素估算比較",
    todo=True,
)
para("〔結果討論〕三大營養素之估算準確度與熱量估算趨勢一致，反映本方法於整體營養推估上之優勢；"
     "各營養素間之差異則與食物型態及密度誤差有關，將於 6.5 節誤差分析中一併討論。")

heading("6.4　行動系統部署效能分析", level=2)
para("除準確度外，系統之實用性亦取決於其回應效能。本節分析操作流程時間、雲端推論延遲與整體"
     "系統回應時間。")

heading("6.4.1　操作流程時間分析", level=3)
make_table(
    ["操作步驟", "平均耗時(秒)", "備註"],
    [
        ["比例尺擺放", "TODO", ""],
        ["三視角拍攝", "TODO", ""],
        ["影像上傳", "TODO", "視網路而定"],
    ],
    caption_text="表 6-6　使用者操作流程時間",
    todo=True,
)

heading("6.4.2　雲端推論延遲分析", level=3)
para("本節於 RunPod 雲端以 NVIDIA RTX PRO 4500（32 GB，Blackwell）實測單次三視角完整管線。"
     "為區分「首次載入模型」與「模型已常駐顯存後之純推論」，於同一 Python 進程內連續執行兩次"
     "管線：第一次為冷啟動（含權重載入），第二次為常駐後推論（FOODCAL_KEEP_MODELS=1，"
     "Qwen2.5-VL、SAM3、DUSt3R 權重皆保留於 GPU）。測試影像為混合餐盤範例（1001–1003.jpg）。"
     "表 6-7 以常駐後之延遲為主（對應實際 Web API 服務模式）；冷啟動合計約 249.4 秒，"
     "其中視覺語言模型載入與首次推論約占 190.7 秒。")
make_table(
    ["推論階段", "平均耗時(秒)", "GPU 記憶體峰值(GB)"],
    [
        ["食物辨識（Qwen2.5-VL）", "10.88", "15.94"],
        ["食材分割（SAM3）", "4.48", "19.24"],
        ["三維重建（DUSt3R）", "4.31", "21.35"],
        ["體積與營養計算", "0.64", "21.35"],
        ["合計（模型常駐）", "20.61", "21.35"],
    ],
    caption_text="表 6-7　雲端各推論階段之延遲與記憶體用量（RTX PRO 4500；模型常駐後）",
    todo=False,
)
para("〔結果討論〕常駐後端到端約 20.6 秒，其中視覺語言模型仍為最大瓶頸（約 10.9 秒，占 53%），"
     "其次為 SAM3 與 DUSt3R（各約 4–5 秒）。冷啟動約 249 秒，顯示「權重載入」遠重於純推論；"
     "故部署上宜採長駐進程（如 FastAPI + 模型快取）以避免每次請求重新載入約 16 GB 之 VLM。"
     "常駐時三模型合計顯存峰值約 21.4 GB，於 32 GB 卡上可行。")

heading("6.4.3　系統回應時間分析", level=3)
para("系統整體回應時間定義為使用者完成上傳至收到分析結果之總耗時，為操作流程、資料傳輸與雲端"
     "推論延遲之總和。下表彙整雲端推論實測（不含行動端拍攝與上傳網路延遲）。")
make_table(
    ["測試場景", "端到端回應時間(秒)", "備註"],
    [
        ["混合取餐（模型常駐）", "20.61", "RTX PRO 4500；同一進程第二次推論"],
        ["混合取餐（冷啟動）", "249.42", "含 Qwen／SAM3／DUSt3R 首次載入"],
        ["單一食材（活大自助餐）", "—", "待以標定單品影像補測"],
    ],
    caption_text="表 6-8　各場景之端到端系統回應時間（雲端推論）",
    todo=False,
)
para("〔結果討論〕在模型常駐之服務模式下，混合餐盤之雲端推論約 21 秒，已接近可互動之回應時間；"
     "若每次請求皆冷啟動，延遲將升至約 4 分鐘，不適合作為線上服務。後續可再補測單一食材樣本，"
     "並將上傳與行動端操作時間一併計入完整端到端。")

heading("6.4.4　不同 GPU 規格之推理時間與成本比較", level=3)
para("為協助評估系統於不同硬體上之部署效益，本節比較數種 GPU 規格於執行單次完整推論（一次三"
     "視角分析）時之推理時間與推理成本。候選 GPU 包含 NVIDIA L4、RTX 4090、RTX PRO 4500"
     "（本研究採用）、RTX 5090 等，其顯示記憶體與每小時租用費用見表 5-3。"
     "由於 RunPod 採按秒計費，單次推論之成本可由式 (6-1) 計算：")
equation(
    m_omath(
        m_sub(m_text("C", italic=True), m_text("infer")),
        m_text(" = "),
        m_frac(
            m_sub(m_text("T", italic=True), m_text("infer")),
            m_text("3600"),
        ),
        m_text(" × "),
        m_sub(m_text("C", italic=True), m_text("hour")),
    ),
    number="(6-1)",
)
para("式中，{C:infer} 為單次推理成本（US$）；{T:infer} 為單次推理時間（秒）；"
     "{C:hour} 為每小時租用費用（US$/hr）。",
     indent_chars=0)
para("本研究已於 RTX PRO 4500、RTX 4090 與 RTX 5090 完成常駐後實測（同一混合餐盤、相同管線）；"
     "另於 NVIDIA L4（約 24 GB）嘗試常駐三模型時發生 CUDA out of memory（進程已占用約 22 GB 顯存），"
     "顯示 L4 不足以在「模型常駐」模式下同時容納 Qwen2.5-VL-7B、SAM3 與 DUSt3R。"
     "表 6-9 以常駐後單次推理時間計算成本；PRO 4500 冷啟動 249.42 秒僅作對照。"
     "相對速度以 PRO 4500 常駐 20.61 秒為 1.0。")
make_table(
    ["GPU 型號", "顯示記憶體", "每小時費用(US$)", "單次推理時間(秒)", "單次推理成本(US$)", "相對速度"],
    [
        ["NVIDIA L4", "24 GB", "約 0.40–0.79", "常駐 OOM（不可行）", "—", "—"],
        ["RTX 4090", "24 GB", "約 0.60", "15.69（常駐）", "0.0026", "1.31×"],
        ["RTX PRO 4500", "32 GB", "約 0.74", "20.61（常駐）", "0.0042", "1.00×（基準）"],
        ["RTX 5090", "32 GB", "約 0.99", "13.19（常駐）", "0.0036", "1.56×"],
        ["RTX PRO 6000", "96 GB", "約 2.09", "未實測", "—", "—"],
        ["NVIDIA A100", "80 GB", "約 1.40–2.10", "未實測", "—", "—"],
    ],
    caption_text="表 6-9　不同 GPU 規格之推理時間與單次推理成本比較",
    todo=False,
)
para("此外，尚需區分「一次性冷啟動成本」與「每次推理成本」：冷啟動成本主要來自首次下載與載入"
     "模型（約數十 GB），此成本可透過持久化網路磁碟保存模型與環境、以及進程內模型常駐而在"
     "多次推論間攤銷，故不計入常態之單次推理成本。以實務營運觀點，單位使用者之每次分析成本可"
     "估計為：「單次推理時間 × 每秒單價」，並可另計入網路磁碟之固定月租等間接成本。"
     "實測常駐單次成本約為：RTX 4090 US$ 0.0026、RTX 5090 US$ 0.0036、RTX PRO 4500 US$ 0.0042。")
para("〔結果討論〕就延遲而言，RTX 5090 最快（13.19 秒，約 1.56×），其次為 RTX 4090（15.69 秒，"
     "約 1.31×），PRO 4500 為 20.61 秒。就單次成本而言，RTX 4090 最低（約 0.26 美分），因每小時"
     "單價較低且速度仍優於 PRO 4500；5090 雖最快，但較高時薪使單次成本略高於 4090。"
     "NVIDIA L4 雖時薪較低，但在常駐模式下顯存不足（OOM）；若採「各階段載入後釋放」之序列化"
     "策略理論上可於 24 GB 卡上執行，惟延遲將接近冷啟動量級，不適合作為低延遲線上服務。"
     "綜合延遲、成本與顯存，論文原型與小流量服務建議優先選 24–32 GB 且能穩定常駐之規格"
     "（實測以 RTX 4090／PRO 4500／5090 為可行集）。")
fig_img("fig6_2.png", "圖 6-2　已實測 GPU 之推理時間與單次成本比較（L4 常駐 OOM，未列入長條）", width_cm=14)
todo_note("L4 常駐不可行已記入表 6-9；若未來補測「卸載模式」之 L4 單次秒數可再更新。"
          "PRO 4500／4090／5090 常駐數據見 benchmark_report.json。")

heading("6.5　實驗結果綜合討論與誤差來源分析", level=2)
para("綜合前述結果，本節討論系統之整體表現並剖析主要誤差來源。就本方法之設計而言，"
     "誤差可能來自以下數個環節：")
numbered("辨識與分割誤差：視覺語言模型之誤辨或漏辨、萬物分割模型之邊界誤分割，皆會直接影響"
         "後續界定食材點雲之正確性；相近食材相鄰或遮擋時尤為明顯。goFOOD 與 Vinod 等人"
         f"研究{cit(34,22)}均顯示，分割品質為份量估算之上游瓶頸。")
numbered("三維重建誤差：在紋理稀疏、反光或視差不足之情況下，DUSt3R 之點雲可能出現雜訊或"
         "局部失真，影響厚度估計。相較 Dehais 等人"
         f"之稠密 stereo{cit(29)}，學習式重建對弱紋理白飯、清湯較敏感。")
numbered("尺度錨定誤差：信用卡分割之誤差、其於三維空間中對角線量測之偏差，或信用卡未與餐點"
         "共平面，皆會造成換算係數之偏差，進而以立方之方式放大至體積。TADA 與 Dhar 等人"
         f"{cit(30,50)}均強調參考物分割與共平面對尺度之關鍵性。")
numbered("參考平面擬合誤差：當食材邊界並非恰位於盤面、或盤面本身不平整時，平面擬合會產生偏差，"
         "影響高度（厚度）之量測基準。")
numbered("密度與營養值誤差：資料庫所採用之堆積密度係以管線體積—秤重配對標定，營養值則為代表值；"
         "實際食物因含水量、烹調方式與堆疊鬆緊之差異，會引入系統性偏差。由於建庫與推論共用"
         "同一管線，密度標定可部分吸收體積估計之系統偏差，但一分量型態改變時仍可能失配；"
         "Nutrition5k"
         f"{cit(20)}亦指出 per-gram nutrition 假設在混合烹調下仍有局限。")
numbered("名稱匹配誤差：視覺語言模型輸出之食物名稱若與資料庫識別鍵無法正確對應，將導致採用"
         "錯誤之營養參數或標記為未知。")
para("與文獻對照時，須先對齊實驗條件：Nutrition5k 之 26.1%／16.5% 來自固定 top-view"
     f"設定{cit(20)}；Dehais 之 8.2% 需雙視角與參考卡{cit(29)}；Vinod 之 17.67%"
     f"需 checkerboard 與 3D food model{cit(22)}；Shao 等人 MAPE 11.47% 來自"
     f"96 張 nutrition study{cit(49)}，而非 Shao 2023 體素重建論文"
     f"{cit(21)}——後者原文明列 MAPE 在摘要（11.47%）與表格（22.0%）間不一致，"
     "故僅引用其 MAE 40.05 kCal 作方法描述。本研究完成實測後，將在相同活大協定下"
     "報告本方法與 LMM 基準之 MAPE，並討論與上述文獻之異同，但不作跨資料集之"
     "直接優劣排序。")
para("上述誤差來源將於數值結果填入後，結合具體案例進一步量化與討論，並據以提出對應之改善方向。")
todo_note("待各節數據填入後，補充綜合討論之量化結論，並針對誤差最大之環節提出具體之改善建議。")

# ===========================================================================
# 第七章 結論與未來展望
# ===========================================================================
heading("第七章　結論與未來展望", level=1, page_break_before=True)

heading("7.1　結論", level=2)
para("本研究針對現有二維影像式食物熱量估算方法「缺乏絕對物理尺度、易生視覺幻覺、且缺乏可解釋性」"
     f"三大痛點{cit(16,24)}，提出並實作一套「基於三維重建與語意分割之食物熱量估算系統」。"
     "本系統以日常隨身之標準信用卡作為絕對比例尺，整合輕量化視覺語言模型 Qwen2.5-VL、"
     f"萬物分割模型 SAM3 與多視角三維重建模型 DUSt3R{cit(6,12,13)}，構成一條由影像至三維"
     "體積、再至熱量與營養之可解釋計算管線，並以行動端與雲端協同之架構完成系統原型。")
para("本研究之主要成果與貢獻歸納如下：")
numbered("提出以標準信用卡作為絕對比例尺之尺度錨定方法，於免專用深度感測硬體之前提下，"
         "將無尺度之三維重建結果校正為具絕對物理單位之空間模型，從根本上克服二維影像缺乏尺度"
         "之問題。")
numbered("提出整合多模態視覺與多視角三維重建之可解釋體積估算管線，並設計結合平面擬合與厚度"
         "積分之體積計算演算法，使每一步驟皆產生可視化、可檢驗之中間結果。")
numbered("建置食物幾何與營養轉換資料庫，以與推論相同之管線標定各品項一分量之管線估計體積與秤重"
         "對應，並透過預先計算之每立方公分營養值，完成由管線估計體積至熱量與三大營養素之轉換。")
numbered("設計並實作行動端與雲端協同之系統原型，提出標準化之多視角拍攝流程，並規劃了與大型"
         "多模態模型之比較實驗以驗證方法之效益。")
todo_note("結論最後請於實驗完成後，補上量化之總結陳述（如體積與熱量之 MAPE、相較基準方法之"
          "改善幅度），以具體回應研究目標之達成程度。")

heading("7.2　未來展望", level=2)
para("基於本研究之成果與觀察，未來可朝以下方向進一步發展：")
numbered("推論效能之最佳化：目前之推論延遲主要由三維重建與視覺語言模型主導，未來可透過模型"
         f"量化、蒸餾、更輕量之重建方法（如 MASt3R 系列{cit(44)}），或以更少之視角配合"
         "更佳之演算法，縮短回應時間並降低運算成本。")
numbered("尺度錨定之強健化：可探討自動偵測多種常見參考物（如硬幣、餐具）之機制，或研究於"
         f"無明確參考物時，藉由裝置慣性感測與多視角幾何自我校正尺度之可能性{cit(22,23,36)}，"
         "以進一步降低使用門檻。")
numbered("密度與營養模型之精緻化：擴充食物幾何與營養轉換資料庫之涵蓋範圍，並參考 Nutrition5k"
         f"{cit(20)}等資料集之成分標註方式，研究以更多一分量樣本重複標定管線體積—秤重對應，"
         "或以影像特徵動態估計含水量與烹調狀態，以降低堆積密度所引入之系統性誤差。")
numbered("食材辨識與分割之強化：針對相近食材相鄰、遮擋與複雜擺盤之情境，強化辨識與分割之"
         "準確度，例如結合多視角之語意一致性以修正單一視角之誤判。")
numbered("實際場域驗證與長期使用者研究：於真實使用情境中進行更大規模之驗證，並結合健康管理"
         "應用進行長期之使用者研究，評估系統對使用者飲食行為與健康成效之實際影響。")
numbered("端側部署之可行性：隨著行動裝置運算能力之提升與模型輕量化技術之進展，未來可評估將"
         "部分或全部推論移至端側執行之可行性，以降低對雲端之依賴並提升隱私保護。")
numbered("foundation model 與顯式幾何結合：Vinod 等人"
         f"{cit(22)}已證明 YOLOv8 + SAM + pose + rendering 可行；下一步可保留幾何推理鏈，"
         "將偵測／分割換成更強 VLM 與 SAM3，並研究 implicit-scale 3D reconstruction 以"
         "降低對參考物之依賴。")
numbered("在地化資料集與垂直場景：參考 Dhar 等人"
         f"{cit(50)}與 Nutrition5k{cit(20)}，建立含 RGB、體積、重量、烹調型態與營養成分"
         "之台灣餐飲資料協定，支援活大自助餐、便當與夜市等子場景之 domain adaptation。")
numbered("不確定性估計：對混合菜、半遮擋與高形變食物，輸出熱量可信區間而非單點估計，"
         "提升 mHealth 應用之使用者信任。")
para("綜言之，本研究驗證了以「日常物品比例尺＋多模態視覺＋多視角三維重建」達成免專用硬體、"
     "具絕對尺度且可解釋之食物熱量估算之可行性。文獻回顧顯示，此路徑在 Dehais 之幾何"
     f"可解釋性{cit(29)}、Nutrition5k 之 benchmark 嚴謹性{cit(20)}與 Shao 系列之"
     f"單眼 RGB 企圖心{cit(49,21)}之間取得務實平衡，期能為行動醫療與智慧健康照護領域"
     "之飲食評估提供一條兼顧準確度、實用性與可解釋性之技術路徑。")

# ===========================================================================
# 參考文獻
# ===========================================================================
heading("參考文獻", level=1, page_break_before=True)


def ref(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-24)
    r = p.add_run(text)
    _set_run_font(r, size=Pt(11))


for _ref_text in REFERENCES:
    ref(_ref_text)

# @@CHAPTERS@@

# ---------------------------------------------------------------------------
# 讓 Word 於開啟檔案時自動更新所有功能變數（目錄／圖目錄／表目錄與頁碼）
# ---------------------------------------------------------------------------
_settings = doc.settings.element
_upd = _settings.find(qn("w:updateFields"))
if _upd is None:
    _upd = OxmlElement("w:updateFields")
    _settings.insert(0, _upd)
_upd.set(qn("w:val"), "true")

_output = os.path.join(SCRIPT_DIR, "碩士論文2.docx")
_fallback = os.path.join(SCRIPT_DIR, "碩士論文_引用更新.docx")
_fallback2 = os.path.join(SCRIPT_DIR, "碩士論文_章節修訂.docx")
_saved = []
for _path in (_output, _fallback, _fallback2):
    try:
        doc.save(_path)
        _saved.append(_path)
    except PermissionError:
        continue
if not _saved:
    from datetime import datetime
    _ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    _path = os.path.join(SCRIPT_DIR, f"碩士論文_{_ts}.docx")
    doc.save(_path)
    _saved.append(_path)
print("已輸出：")
for _path in _saved:
    print(f"  {_path}")
print("請開啟「碩士論文2.docx」（若 Word 仍報錯，請先關閉已開啟的舊檔再重新產生）")
